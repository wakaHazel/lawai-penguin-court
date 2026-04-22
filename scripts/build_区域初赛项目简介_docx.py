from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
DEFAULT_OUTPUT_PATH = OUTPUT_DIR / "区域初赛提交材料_项目简介_企鹅法庭_论文标准版.docx"
OUTPUT_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUTPUT_PATH


def set_run_font(run, east_asia: str, size: int, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)


def clear_paragraph(paragraph) -> None:
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)


def _append_field_run(paragraph, instr: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run_begin = paragraph.add_run()
    run_begin._r.append(begin)
    run_instr = paragraph.add_run()
    run_instr._r.append(instr_text)
    run_sep = paragraph.add_run()
    run_sep._r.append(separate)
    run_text = paragraph.add_run("1")
    run_text._r.append(end)
    set_run_font(run_text, "宋体", 9, color=RGBColor(0x66, 0x66, 0x66))


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False, size: int = 10, font: str = "宋体") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title_paragraph(doc: Document, text: str, size: int, *, bold: bool = True, spacing_after: int = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    if level == 1:
        paragraph.style = doc.styles["Heading 1"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
    else:
        paragraph.style = doc.styles["Heading 2"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", 16 if level == 1 else 14, bold=True)


def add_body(doc: Document, text: str, *, indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.85)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12)


def add_list_item(doc: Document, prefix: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.hanging_indent = Cm(0.5)
    paragraph.paragraph_format.space_after = Pt(4)
    run_prefix = paragraph.add_run(prefix)
    set_run_font(run_prefix, "宋体", 12, bold=True)
    run_text = paragraph.add_run(text)
    set_run_font(run_text, "宋体", 12)


def add_catalog_entry(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True, center=True, font="黑体")
        clear_cell_shading(header_cells[idx])
        if col_widths_cm:
            header_cells[idx].width = Cm(col_widths_cm[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=10)
            if col_widths_cm:
                cells[idx].width = Cm(col_widths_cm[idx])
    doc.add_paragraph()


def add_table_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 10, bold=False)


def add_reference_item(doc: Document, index: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.hanging_indent = Cm(0.74)
    paragraph.paragraph_format.space_after = Pt(4)
    run_prefix = paragraph.add_run(f"[{index}] ")
    set_run_font(run_prefix, "宋体", 11, bold=True)
    run_text = paragraph.add_run(text)
    set_run_font(run_text, "宋体", 11)


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def configure_section_layout(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.6)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)
    section.different_first_page_header_footer = False


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_section_header(section, text: str | None = None) -> None:
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, "宋体", 10, color=RGBColor(0x44, 0x44, 0x44))


def set_section_footer(section, *, roman: bool = False, visible: bool = True) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not visible:
        return
    prefix = paragraph.add_run("— ")
    set_run_font(prefix, "宋体", 9, color=RGBColor(0x66, 0x66, 0x66))
    _append_field_run(paragraph, "PAGE \\* ROMAN" if roman else "PAGE")
    suffix = paragraph.add_run(" —")
    set_run_font(suffix, "宋体", 9, color=RGBColor(0x66, 0x66, 0x66))


def add_cover_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, center=True, size=11, font="黑体")
        set_cell_text(cells[1], value, size=11)
        clear_cell_shading(cells[0])
        clear_cell_shading(cells[1])
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(10.4)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    configure_section_layout(section)

    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(12)
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    set_section_header(section, None)
    set_section_footer(section, visible=False)


def build_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    add_title_paragraph(doc, "【D06】法律AI应用创新与实践", 18, spacing_after=18)
    add_title_paragraph(doc, "企鹅法庭·沉浸式庭审模拟与法律辅助分析系统", 24, spacing_after=12)
    add_title_paragraph(doc, "区域初赛项目研究报告", 16, spacing_after=16)
    add_title_paragraph(doc, "基于腾讯系AI工具链的智慧法律系统设计与应用研究", 13, bold=False, spacing_after=28)
    for _ in range(3):
        doc.add_paragraph()
    add_cover_meta_table(
        doc,
        [
            ("赛题名称", "【D06】法律AI应用创新与实践"),
            ("赛题方向", "智慧法律应用类"),
            ("项目名称", "企鹅法庭·沉浸式庭审模拟与法律辅助分析系统"),
            ("作品类型", "系统解决方案"),
            ("项目团队", "企鹅法庭项目组"),
            ("完成时间", "2026年4月"),
        ],
    )
    add_title_paragraph(doc, "提交阶段：区域初赛", 12, bold=False, spacing_after=8)
    add_title_paragraph(doc, "文稿属性：正式研究报告", 12, bold=False, spacing_after=8)


def build_catalog(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "目  录", level=1)
    catalog_items = [
        "第一章  引言",
        "第二章  项目基础信息与总体定位",
        "第三章  项目需求分析与问题建模",
        "第四章  系统设计与关键功能实现",
        "第五章  项目使用场景与演示验证",
        "第六章  AI工具链、工作流与Prompt设计",
        "第七章  项目价值、创新点与实施展望",
        "参考文献",
        "附录  区域初赛材料映射表",
    ]
    for item in catalog_items:
        add_catalog_entry(doc, item)


def build_abstract(doc: Document) -> None:
    add_heading(doc, "摘  要", level=1)
    add_body(
        doc,
        "随着法律服务数字化进程的持续推进，法律行业对于能够兼顾专业性、可解释性与可交互性的智慧法律应用提出了更高要求。针对普通诉讼参与人难以在庭前有效整理案情、预判对方抗辩和评估证据强弱的问题，项目围绕【D06】法律AI应用创新与实践赛题构建了“企鹅法庭·沉浸式庭审模拟与法律辅助分析系统”。该系统以民事诉讼庭前准备场景为核心，构建了“案件录入—庭审模拟—对方行为推演—胜诉率分析—复盘报告”的完整主轴，在产品形态上突破了传统法律问答工具停留于单点回答的局限。",
        indent=False,
    )
    add_body(
        doc,
        "在技术方案上，项目采用“前端交互层—后端控制层—腾讯元器工作流编排层—法律能力工具层”的组合架构，以腾讯元器承担流程编排与智能体协同，以得理开放平台API和小理AI承担法律检索与专业支撑，以后端状态控制保证多阶段庭审模拟的稳定推进。在应用设计上，系统将民事诉讼典型流程转化为分阶段、可分支、可复盘的文游式互动过程，使用户能够在接近真实法庭语境的环境中进行策略选择、对抗预判和准备校正。从当前系统实现、工作流设计与演示链路看，项目在普通当事人诉前准备、法学生模拟法庭训练以及智慧法律展示等场景中已经形成较强的适配性和展示潜力，也为法律AI系统从“问答型”向“过程型、策略型、结果沉淀型”演进提供了可落地的系统样本。",
        indent=False,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run_label = paragraph.add_run("关键词：")
    set_run_font(run_label, "黑体", 12, bold=True)
    run_text = paragraph.add_run("智慧法律；庭审模拟；诉前准备；腾讯元器；法律AI；系统解决方案")
    set_run_font(run_text, "宋体", 12)


def build_introduction(doc: Document) -> None:
    add_heading(doc, "第一章  引言", level=1)
    add_body(
        doc,
        "在智慧法律应用快速发展的背景下，法律AI系统正在从单纯的知识问答、法条检索和文书生成，逐步走向更强调过程协助、情境推演和结果沉淀的综合型系统。然而，从现实法律服务场景来看，用户在面对具体案件时最迫切的困难，往往不是完全找不到法条，而是不知道如何围绕争议焦点组织事实、如何判断证据链是否完整、如何预判对方抗辩，以及如何在正式庭审前形成一套可执行的准备方案。现有多数工具虽然能够提供局部帮助，但在“把案件真正走一遍”这一层面上仍存在明显不足。",
    )
    add_body(
        doc,
        "基于此，本项目选择“庭前准备与模拟推演”作为系统切入点，以普通诉讼参与人、法学生模拟法庭训练者以及智慧法律展示场景为主要对象，提出并实现“企鹅法庭·沉浸式庭审模拟与法律辅助分析系统”。项目强调通过流程化、结构化与互动化的产品设计，将案件整理、程序推进、对方行为预判、胜诉率分析和复盘报告整合为一个完整闭环，以此回应法律场景中“需要实战辅助而非泛化回答”的核心问题。",
    )
    add_heading(doc, "（一）研究背景", level=2)
    add_body(
        doc,
        "赛题《法律AI应用创新与实践》明确鼓励参赛团队依托腾讯系AI工具平台，在法律专业场景中构建具备法律咨询、案件分析、合规审查或系统解决方案能力的作品。其核心导向在于围绕真实法律需求，完成从需求分析、技术方案设计、工作流搭建、核心功能开发到专业验证与演示交付的完整链路。具有竞争力的项目既要满足工具使用约束，也要对法律场景中的真实痛点形成扎实回应。",
    )
    add_heading(doc, "（二）问题界定", level=2)
    add_body(
        doc,
        "核心问题是：如何在赛题允许的腾讯系AI工具链框架下，构建一套既具有法律场景真实性、又具备演示完整性与工程稳定性的庭审模拟系统，使其能够帮助用户在庭前完成案情结构化、对抗预判、证据缺口识别和诉讼准备沉淀。与传统法律问答型产品不同，该问题更强调“过程”和“策略”，也更适合通过工作流与状态控制来组织系统能力。",
    )
    add_heading(doc, "（三）研究思路与报告结构", level=2)
    add_body(
        doc,
        "报告内容依次覆盖项目基础信息与总体定位、用户需求分析与问题建模、系统设计、关键功能、应用场景、AI工具链与工作流实现，以及项目价值、创新点与实施展望。",
    )


def build_section_basic_info(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "第二章  项目基础信息与总体定位", level=1)
    add_body(
        doc,
        "“企鹅法庭”项目面向【D06】法律AI应用创新与实践赛题中的智慧法律应用方向，定位为一套围绕民事诉讼准备与庭审推演构建的沉浸式法律辅助分析系统。项目聚焦的不是单一法律问答，而是把当事人、法学生或法律工作者在“上庭前究竟该准备什么、如何判断自己的证据与论证是否站得住、如何预判对方可能的抗辩路径”这一系列高强度、强流程、强专业的问题，整合进一个可操作、可交互、可复盘的完整系统之中。",
    )
    add_body(
        doc,
        "与传统法律服务工具常见的“检索一个法条、生成一段回答、输出一份文书模板”的模式不同，本项目的主轴被定义为“案件录入—庭审模拟—对方行为推演—胜诉率分析—复盘报告”。系统不会停留在单点回答层面，而是通过结构化案件输入、固定阶段推进、对方行为模拟和结果归纳，使用户能够以更接近真实庭审的方式参与案件分析过程，最终获得一份可执行的诉前准备结果。",
    )

    add_heading(doc, "（一）项目基本属性", level=2)
    add_table_title(doc, "表2-1  项目基本属性")
    add_table(
        doc,
        ["项目字段", "具体内容"],
        [
            ["项目名称", "企鹅法庭·沉浸式庭审模拟与法律辅助分析系统"],
            ["赛题名称", "【D06】法律AI应用创新与实践"],
            ["赛题方向", "智慧法律应用类"],
            ["作品类型", "系统解决方案"],
            ["项目定位", "基于腾讯系AI工具平台构建的沉浸式庭审模拟与法律辅助分析系统"],
            ["核心主轴", "案件录入 -> 庭审模拟 -> 对方行为推演 -> 胜诉率分析 -> 复盘报告"],
            ["目标对象", "普通诉讼参与人、法学生、模拟法庭训练者、智慧法律应用展示场景"],
            ["官方工具链", "腾讯开悟比赛平台、CodeBuddy、腾讯元器、腾讯元宝、小理AI、得理开放平台API"],
        ],
        col_widths_cm=[4.2, 11.2],
    )

    add_heading(doc, "（二）项目建设背景", level=2)
    add_body(
        doc,
        "当前法律服务领域一方面面临案件数量增长、法律服务资源分布不均、普通当事人专业能力不足等现实问题，另一方面也迎来了大模型与智能体技术在法律文本理解、类案检索、争点归纳和文书生成方面的快速发展窗口。赛题本身已经明确指出，智慧法律的关键不只是“会说法律术语”，而是要把 AI 能力真正嵌入法律场景，提高司法辅助效率、降低法律服务门槛，并增强法律服务的可及性和可解释性。",
    )
    add_body(
        doc,
        "在这种背景下，团队选择“模拟法庭备战”作为切入点，原因在于这一场景同时具备几个鲜明特点：第一，场景真实且刚需，无论是普通用户准备诉讼，还是法学生参加模拟法庭训练，都需要提前整理案情、梳理论点和预判对抗；第二，流程清晰且可结构化，民事诉讼法规定的庭审阶段天然适合做工作流和状态机控制；第三，能够充分体现 AI 在法律场景中的综合能力，不仅需要理解案件事实，还需要调动检索、推理、表达、生成和风险提示等能力。因此，“企鹅法庭”并非泛化型产品，而是围绕一条高价值主线，去做深做实的场景化系统。",
    )

    add_heading(doc, "（三）项目建设目标", level=2)
    add_body(
        doc,
        "本项目的建设目标是形成一套真正能够帮助用户进行诉前准备和模拟推演的法律 AI 系统。在产品层面，系统要让用户感受到明确的流程推进和互动反馈，而不是被动接受一整段模型输出；在专业层面，系统要尽可能对齐真实庭审语境，能够围绕争议焦点、证据链完整性、程序性风险和对方抗辩逻辑给出具体判断；在比赛层面，系统要满足赛题对腾讯系工具使用、工作流与 Prompt 说明、专业验证测试和完整演示闭环的要求，形成一套可展示、可答辩、可复现的智慧法律应用方案。",
    )
    add_body(
        doc,
        "项目建设结果包括三个层面：一是帮助普通用户完成“看懂自己的案子”，把零散事实、证据和诉求变成结构化案件输入；二是帮助用户完成“演练自己的庭审”，通过分阶段、多分支的文游式互动感知对抗风险和策略差异；三是帮助用户完成“整理自己的准备”，把模拟结果沉淀为风险评估、胜诉率分析、证据缺口提示和复盘报告。三个层面共同构成“输入—推演—输出”的完整闭环。",
    )

    add_heading(doc, "（四）项目总体架构与系统形态", level=2)
    add_body(
        doc,
        "从系统形态上看，“企鹅法庭”采用的是“前端交互层 + 后端控制层 + 元器工作流编排层 + 法律能力工具层”的组合架构。前端负责案件录入、庭审叙事展示、分支动作点击、结果面板和报告查看；后端负责案件会话管理、模拟状态持久化、结果聚合、兜底逻辑和历史记录管理；腾讯元器承担智能体工作流编排职责，负责案件识别、法律检索、庭审阶段内容生成、对方行为模拟和结果聚合；得理开放平台 API 与小理AI则提供法律法规、类案与法律表达参考的外部支撑。",
    )
    add_body(
        doc,
        "之所以采用这样的结构，是因为赛题既要求官方工具链深度参与，也要求作品具备工程可运行性和展示稳定性。若完全依赖自由生成式对话，系统很难保证庭审阶段不混乱、输出格式不漂移；若完全用静态模板，又无法体现 AI 在法律分析和场景交互上的价值。因此，项目选择“流程优先、生成受控、后端兜底”的方案，用工作流和状态机制约模型边界，用后端保证结果的落地和稳定，从而在专业性、可玩性和演示可控性之间取得平衡。",
    )

    add_heading(doc, "（五）项目创新点概述", level=2)
    add_body(
        doc,
        "本项目的创新点主要体现在三个方面。第一，交互创新：系统不是传统聊天窗口，而是以庭审阶段为核心组织页面和互动逻辑，让用户像经历一次文字冒险式法庭演练一样完成整条诉讼准备链路。第二，流程创新：把法律问答、法条类案检索、对方行为预判、风险评分与报告输出整合进一条有明确顺序、有明确节点、有明确结果沉淀的主流程中。第三，落地创新：在赛题要求的腾讯系工具链基础上，将元器工作流与后端控制层协同使用，既满足官方工具使用要求，又提升了产品稳定性和可演示性。",
    )
    add_body(
        doc,
        "这种创新不是停留在概念层，而是直接指向实际法律应用中的真实问题。普通当事人面临的关键困难往往不是能否搜到法条，而是上庭时是否会被问住、证据是否足够、对方可能如何拆解己方说法。“企鹅法庭”把这些过去主要依靠经验、难以提前练习的内容，转化为一套可以在赛题框架下实现、也具备进一步落地空间的法律 AI 系统。",
    )


def build_section_needs(doc: Document) -> None:
    add_heading(doc, "第三章  项目需求分析与问题建模", level=1)
    add_body(
        doc,
        "需求分析是本项目成立的核心基础。若只是为了展示 AI 可以“写一段法律回答”，完全没有必要构建“企鹅法庭”这样一套系统。团队在前期调研和方案推敲中发现，无论是普通当事人、法学生，还是在法律教学与普法展示中的使用者，他们真正缺的不是某一个孤立功能，而是缺少一种能够把案件事实、法律规则和庭审过程串起来的“实战化辅助工具”。项目的需求分析，也因此围绕用户痛点、场景需求和赛题要求三个维度展开。",
    )

    add_heading(doc, "（一）行业与场景背景分析", level=2)
    add_body(
        doc,
        "在传统法律服务链条中，普通当事人在准备诉讼时经常面临信息断裂：案件事实掌握在自己手里，但不知道哪些信息是法律上真正重要的；手中持有证据材料，但不知道哪些能证明关键事实、哪些证据会在法庭上被质疑；能够通过搜索获得一些法律知识，但无法把法条、案例与自己案件中的争议点有效对应起来。最终导致的结果往往是，用户在正式庭审前既缺乏清晰的整体认知，也缺乏具体的准备方案。",
    )
    add_body(
        doc,
        "与此同时，在法学教育与模拟法庭训练场景中，也存在另外一种典型问题：学生知道法庭程序、知道部分法条，但对真实攻防节奏缺乏体感，对“对方会提出什么抗辩”“法官在这个节点会关注什么”“自己哪里会被追问”缺少系统认知。很多训练仍停留在静态材料阅读和结果复述层面，缺乏能够让学生在过程里做决策、在对抗中感受后果的沉浸式工具。也正因为如此，面向庭审模拟场景的智慧法律系统兼具现实价值与教学价值。",
    )

    add_heading(doc, "（二）目标用户分析", level=2)
    add_table_title(doc, "表3-1  目标用户及核心诉求分析")
    add_table(
        doc,
        ["用户类型", "主要特征", "核心诉求", "项目可提供的价值"],
        [
            ["普通诉讼参与人", "缺乏系统法律训练，案情掌握零散，证据判断能力有限", "看懂自己的案子、知道该怎么准备、预判上庭风险", "结构化录入案情、模拟庭审攻防、输出证据缺口与准备建议"],
            ["法学生与模拟法庭队员", "有一定法律知识，但缺乏真实对抗经验", "训练法庭表达、感受庭审流程、验证策略选择差异", "提供阶段化庭审推进、对方行为模拟和复盘报告"],
            ["法律教学与普法展示人群", "需要可视化、可演示、可讲解的智慧法律样例", "用一个系统展示 AI 与法律结合的具体形态", "形成完整主链路、界面交互和报告成果，便于展示与答辩"],
        ],
        col_widths_cm=[2.5, 4.0, 4.0, 4.9],
    )
    add_body(
        doc,
        "从以上分析可以看出，这三类用户虽然使用身份不同，但痛点具有一致性：都希望把抽象的法律知识转化为可操作、可验证、可推进的过程性工具。换言之，他们真正需要的是一套能够“帮助自己进行法庭准备和过程演练”的系统，而不是简单的法律知识检索器或文本生成器。这也是“企鹅法庭”采取流程化产品设计的根本原因。",
    )

    add_heading(doc, "（三）核心问题与痛点拆解", level=2)
    add_table_title(doc, "表3-2  核心痛点与系统需求映射")
    add_table(
        doc,
        ["痛点编号", "典型问题", "用户实际表现", "对系统的需求要求"],
        [
            ["P1", "案件事实难以结构化整理", "用户只能口语化叙述案情，难以抓住关键事实与争点", "系统需要提供分步录入、字段引导和案件摘要整理能力"],
            ["P2", "证据强弱难以判断", "用户不知道哪些证据能证明何种事实，也不知道证据链是否完整", "系统需要对证据进行用途归类、缺口提示和风险提醒"],
            ["P3", "对庭审流程缺乏体感", "用户知道要上庭，但不知道庭上每个阶段会发生什么", "系统需要严格按庭审流程推进并解释当前阶段任务"],
            ["P4", "无法预判对方的抗辩与突袭", "用户往往只准备自己的说法，对对方策略几乎没有准备", "系统需要模拟对方证据、抗辩意见与质证路径"],
            ["P5", "AI建议停留在宏观层面", "获得的建议常常是“完善证据”“注意表达”等宽泛表述", "系统需要输出更具体的证据清单、发言要点、风险点和下一步动作"],
        ],
        col_widths_cm=[1.8, 3.5, 5.0, 5.1],
    )

    add_heading(doc, "（四）赛题需求与项目需求的对应关系", level=2)
    add_body(
        doc,
        "赛题对于智慧法律应用的要求，并不是单纯考察一个大模型是否能回答法律问题，而是强调：要依托腾讯系 AI 工具平台，围绕法律场景开发一个具备专业理解、流程设计、技术实现、专业验证和可演示性的智能体、工具或系统。对照这一要求，“企鹅法庭”的需求设计与赛题目标具有较高一致性。项目围绕法律专业场景的真实需求组织功能模块，并在工作流设计、工程开发、内容优化和法律检索环节形成了腾讯元器、腾讯元宝、CodeBuddy、腾讯开悟平台、小理AI和得理开放平台API的分工关系。",
    )
    add_body(
        doc,
        "赛题要求中的“法律需求分析、技术方案设计、核心功能开发、专业验证测试、迭代优化演示”，在本项目中分别对应：前期对诉讼准备与模拟法庭场景的痛点拆解、系统的工作流与前后端协同设计、案件录入与庭审模拟等核心模块开发、法学同学对法律表述与推理逻辑的人工校验，以及最终形成可交互 demo 和可提交文稿。项目从诉讼准备场景出发，沿着赛题要求所鼓励的路径形成系统化成果。",
    )

    add_heading(doc, "（五）项目需求优先级判断", level=2)
    add_body(
        doc,
        "在需求优先级上，团队明确采取“主链路优先、演示优先、法律真实感优先”的策略。最优先满足的是：案件可以录入、庭审可以推进、对方行为可以模拟、结果可以分析、报告可以生成。这条主链路打通后，项目即具备产品展示和赛题答辩的核心价值。更大范围的案件类型覆盖、更复杂的多人协同、更加完整的立法沙盘拓展等能力，放在后续迭代阶段。",
    )
    add_body(
        doc,
        "这样的优先级判断，并不是对功能的删减，而是对比赛周期和项目主轴的尊重。对于区域初赛来说，评委首先关心的是：项目是否明确解决了一个真实问题，是否形成了闭环，是否体现了工具链使用，是否有足够完成度。基于这一判断，“企鹅法庭”将所有需求都收束到“帮助用户完成一次可感知、可推演、可复盘的庭前准备体验”上，从而使项目在价值表达和成果呈现上更加聚焦。",
    )


def build_section_features(doc: Document) -> None:
    add_heading(doc, "第四章  系统设计与关键功能实现", level=1)
    add_body(
        doc,
        "围绕前述需求分析，“企鹅法庭”最终形成的不是一组松散功能，而是一条以庭审准备为中心的产品主轴。为了保证项目介绍部分既完整又聚焦，本节将按照用户实际使用顺序，对系统的关键功能进行拆解说明。所有功能都服务于同一个目标：帮助用户把案件信息转化为行动准备，把模拟过程转化为结果复盘。",
    )

    add_heading(doc, "（一）案件录入功能", level=2)
    add_body(
        doc,
        "案件录入是整套系统的入口模块，也是后续庭审模拟和结果分析的基础。系统不要求用户一次性输入一大段自由文本，而是通过分步结构化表单引导用户填写案件类型、基本事实、诉讼请求、证据材料、争议焦点和对方信息等内容。这样做的目的是把原本零散、口语化、难以直接计算的案情信息，转换成后续工作流可以识别和利用的结构化输入。",
    )
    add_body(
        doc,
        "在这一模块中，系统不仅承担“收集信息”的职责，还承担“辅助整理信息”的职责。例如，系统会根据用户输入的事实和请求，提炼案件摘要；根据已填信息提示可能存在的争议焦点；根据证据清单识别缺失项和薄弱环节。案件录入本身就是一次轻量级案情梳理过程，不是简单表单，而是用户第一次“把案件看清楚”的过程。",
    )

    add_heading(doc, "（二）沉浸式庭审模拟功能", level=2)
    add_body(
        doc,
        "沉浸式庭审模拟是本项目的核心中的核心。系统严格参照民事诉讼中的典型庭审流程，将模拟过程划分为开庭准备、法庭调查、举证质证、法庭辩论、最后陈述、调解/判决以及复盘报告等阶段。每个阶段都会围绕当前的程序任务、事实焦点和证据状态生成相应叙事，并在关键节点向用户开放策略选择，使用户能够在推进过程中真实感受到“我此刻在法庭上的处境是什么”。",
    )
    add_body(
        doc,
        "与普通聊天式交互不同，本系统采用“分阶段叙事 + 节点动作选择”的模式。用户不是面对无限自由的对话框，而是在系统给出的当前法庭场景下做出更接近真实诉讼策略的选择，例如是否对证据真实性提出异议、是否申请证人出庭、是否要求补充调查、是否在辩论阶段集中攻击对方证明责任不足等。每一次选择都会对后续叙事、对方行为和风险评估产生影响，这也是项目沉浸感和实战感的重要来源。",
    )

    add_heading(doc, "（三）对方行为模拟功能", level=2)
    add_body(
        doc,
        "真实庭审中，决定用户感受和结果走向的，往往不仅仅是自己说了什么，更在于对方会怎么拆解自己的论点、会不会突然提交证据、会从哪个角度质疑自己的证据链。“企鹅法庭”为此专门设计了对方行为模拟能力。系统会依据案件类型、对方身份、当前阶段、已知事实和已有证据，生成对方可能提出的抗辩意见、质证内容、补充证据和策略动作。",
    )
    add_body(
        doc,
        "这一功能的价值在于让用户意识到：庭审不是单向表达，而是一个动态对抗过程。用户在模拟中会看到，对方可能主张事实并非如此、某份证据关联性不足、某项损失计算没有依据，甚至可能在程序节点上主动争取延长举证或要求重新鉴定。通过这种对抗式模拟，用户能够提前意识到自己的薄弱点，也能提前准备应对口径，从而把系统体验从“法律信息阅读”升级为“诉讼策略演练”。",
    )

    add_heading(doc, "（四）胜诉率分析与风险评估功能", level=2)
    add_body(
        doc,
        "在项目设计中，“胜诉率分析”并不被理解为一个绝对预测数字，而被理解为一套围绕案件胜算、证据采信、程序风险和争点覆盖度的综合评估能力。系统会基于案件录入信息、模拟过程中出现的关键节点、用户的策略选择、对方的抗辩表现以及法律检索结果，对案件进行多维度判断，并以更易理解的方式呈现给用户。",
    )
    add_body(
        doc,
        "这一功能之所以重要，是因为许多用户最关心的问题并不是抽象的法律概念，而是“我这案子大概站不站得住”“我现在最大的问题到底是什么”。系统在这部分输出中，不仅会给出整体结论，还会拆分为争点支撑度、证据完整度、程序风险、对方压力值等多个维度，并对每个维度作出解释。这样，用户看到的不再是一个空泛结论，而是一套“问题在哪里、为什么在那里、下一步该怎么办”的可执行分析结果。",
    )

    add_heading(doc, "（五）复盘报告与备战清单功能", level=2)
    add_body(
        doc,
        "复盘报告是系统输出价值最终落地的关键环节。在完成一轮模拟之后，系统会将案件摘要、关键争点、庭审过程中的重要节点、对方可能的攻击路径、己方暴露出的风险点、证据缺口和后续准备建议整合成一份可阅读、可展示、可继续使用的结构化结果文稿。该报告既将互动结果沉淀为“下一步行动清单”，也是系统能够稳定输出成果的重要载体。",
    )
    add_body(
        doc,
        "与复盘报告相配套的，是系统进一步输出的庭前备战清单。该清单会聚焦“用户下一步最该补什么”，例如还缺哪些证据、哪些事实需要进一步固定、哪些说理链条需要补强、哪些法条和类案应当重点关注、面对对方可能的抗辩时应准备怎样的回应。通过这一功能，整个系统真正完成了从案情输入、过程模拟到准备结果输出的闭环，使用户不仅完成了模拟，还获得了可继续推进现实诉讼准备的依据。",
    )

    add_heading(doc, "（六）关键功能一览表", level=2)
    add_table_title(doc, "表4-1  系统关键功能与输出关系")
    add_table(
        doc,
        ["功能模块", "核心作用", "主要输入", "主要输出", "对应价值"],
        [
            ["案件录入", "将案情与证据结构化", "案件事实、请求、证据、对方信息", "案件摘要、争点识别、结构化档案", "帮助用户看清自己的案子"],
            ["庭审模拟", "按法定阶段推进模拟互动", "案件状态、当前阶段、用户动作", "场景叙事、可选策略、阶段变化", "帮助用户感受真实庭审节奏"],
            ["对方行为模拟", "生成对抗式抗辩与质证内容", "案件类型、对方画像、当前焦点", "抗辩意见、证据突袭、程序动作", "帮助用户预判对方思路"],
            ["胜诉率分析", "多维度评估案件胜算与风险", "模拟过程、检索结果、证据状态", "胜算判断、风险点、维度解释", "帮助用户识别关键问题"],
            ["复盘报告/备战清单", "沉淀模拟成果并给出行动建议", "全流程事件与分析结果", "复盘报告、证据清单、准备建议", "帮助用户形成下一步方案"],
        ],
        col_widths_cm=[2.3, 3.0, 3.3, 3.3, 4.0],
    )


def build_section_scenarios(doc: Document) -> None:
    add_heading(doc, "第五章  项目使用场景与演示验证", level=1)
    add_body(
        doc,
        "一个优秀的智慧法律项目，不仅要有功能，更要说清楚“谁会在什么情况下使用它、使用之后能得到什么结果”。“企鹅法庭”的使用场景并非单一，而是围绕诉讼准备、法学训练和智慧法律展示形成了多个可落地的应用方向。不同场景下，系统的切入重点略有差异，但底层价值是一致的，即帮助用户在正式进入法律程序之前，先把问题看清、演练清楚、准备到位。",
    )

    add_heading(doc, "（一）普通当事人的诉前准备场景", level=2)
    add_body(
        doc,
        "普通当事人最常见的问题并不是完全不懂法律，而是不知道如何把自己掌握的事实和证据组织成一个“法庭能听懂、法庭会认可”的表达结构。例如，劳动争议案件中的劳动关系证明、民间借贷案件中的借款事实与资金流向、侵权案件中的因果关系与损失证明，很多时候当事人已有部分材料，但并不知道哪些最关键、缺什么、会被对方怎样拆解。这是“企鹅法庭”直接发挥价值的应用场景。",
    )
    add_body(
        doc,
        "在这一场景中，用户可以先录入自己的案情和现有证据，再通过庭审模拟去观察系统如何归纳争点、法官在不同节点可能会问什么、对方可能从何处发起抗辩，最后通过胜诉率分析和备战清单明确下一步准备重点。系统的意义不在于代替律师作出最终判断，而在于帮助用户把“我大概知道点什么”转变为“我知道我现在最该准备什么”。这对于降低普通人诉讼准备门槛具有明显价值。",
    )

    add_heading(doc, "（二）法学生与模拟法庭训练场景", level=2)
    add_body(
        doc,
        "法学生在学习民事诉讼法、证据法和实体法规则时，往往能够掌握条文和案例，但对真实法庭中的动态攻防、临场表达和策略选择缺乏体验。许多模拟法庭训练也更多依赖线下集中排演，成本较高、频次有限，且难以让每个学生都在同样条件下多次尝试不同策略。“企鹅法庭”恰好能够补足这一空缺，为法学生提供一种低成本、可反复、可分支比较的训练方式。",
    )
    add_body(
        doc,
        "在这一场景下，系统不仅可以用于个人练习，也可以用于课程辅助和团队训练。教师或指导者可以预设案例，学生分别以不同策略推进相同案件，再对比不同分支下的模拟走向与结果报告，从而更直观地理解法庭程序、论证方法和证据组织之间的关系。相比于单纯阅读教材或案例裁判文书，这种带有过程反馈的交互方式，更容易帮助学生形成“法庭语感”和“争点意识”。",
    )

    add_heading(doc, "（三）法律援助、普法与体验展示场景", level=2)
    add_body(
        doc,
        "在法律援助咨询、普法宣传和智慧法律展示场景中，最大的难点通常是如何把专业内容变得可理解、可参与、可感知。很多传统法律服务演示容易停留在“检索一下”“问答一下”的层面，观众难以真正感受到法律 AI 的价值。而“企鹅法庭”由于具有清晰的流程主轴和较强的互动性，非常适合作为智慧法律场景中的展示型应用，用于体现 AI 如何进入法律流程、如何辅助分析、如何输出成果。",
    )
    add_body(
        doc,
        "具体而言，在高校普法活动、法律科技展示、创新创业路演或智慧司法体验活动中，工作人员可以直接选择一个 Demo 案件，让观众看到系统如何从案情录入进入庭审模拟，再如何生成风险分析与备战报告。这样的展示不但更具故事性和可看性，也更容易让非专业观众理解：AI 在法律场景中的价值，不只是“替你搜到资料”，而是“帮助你走完一条本来很难走清楚的准备流程”。",
    )

    add_heading(doc, "（四）比赛展示与答辩场景", level=2)
    add_body(
        doc,
        "对于本次区域初赛及后续答辩而言，项目本身也需要具备良好的“比赛展示场景适配能力”。评委通常会在较短时间内判断一个作品是否真正聚焦问题、是否形成闭环、是否能够清楚说明官方工具链如何参与、是否具有创新表达和完成度。因此，项目必须能够在 5 至 8 分钟的演示窗口中，让评委快速看懂产品价值、看到交互过程、理解技术逻辑，并确认系统已经形成可展示的专业输出链路。",
    )
    add_body(
        doc,
        "“企鹅法庭”的展示方式非常适合这一场景：先用一段简明介绍说明项目定位，再快速进入一个典型案件，展示录入、推进、分支选择、对方行为模拟和报告生成的全过程，最后用复盘报告与 AI 工具说明收尾。这种“从输入到输出”的完整链条，比单点功能展示更有说服力，也更能体现项目作为系统解决方案的完整性。对于比赛场景来说，这种可讲述、可演示、可证明成果的能力，是作品竞争力的重要组成部分。",
    )

    add_heading(doc, "（五）典型使用流程示意", level=2)
    add_table_title(doc, "表5-1  典型应用场景与使用流程")
    add_table(
        doc,
        ["场景类型", "进入方式", "核心交互", "用户获得的结果"],
        [
            ["普通用户诉前准备", "录入真实或拟真实案件信息", "案件整理 -> 庭审模拟 -> 结果分析", "证据缺口、风险提示、备战清单"],
            ["法学生训练", "导入课程案例或训练案例", "多轮分支对比 -> 复盘讨论", "程序理解、论证训练、策略差异感知"],
            ["普法/展示活动", "选择预置 Demo 案件", "演示一轮完整流程", "直观理解法律 AI 应用形态"],
            ["比赛答辩", "使用高完成度 Demo 案件", "聚焦主链路展示与工具链说明", "向评委展示项目闭环与创新点"],
        ],
        col_widths_cm=[3.0, 3.5, 4.5, 4.6],
    )


def build_section_ai_tools(doc: Document) -> None:
    add_heading(doc, "第六章  AI工具链、工作流与Prompt设计", level=1)
    add_body(
        doc,
        "根据赛题要求，本项目在 AI 工具与能力链路上，严格围绕腾讯系官方工具平台与赛题允许使用的法律能力服务展开。为了避免在材料中出现“工具写了很多，但分工不清、工作流不明、Prompt 设计模糊”的问题，本节将从工具职责、工作流主链路、节点设计和 Prompt 设计四个层面，说明项目如何真正使用这些 AI 工具来构建系统，而不是仅仅在文稿中列出工具名称。",
    )

    add_heading(doc, "（一）项目使用的官方 AI 工具", level=2)
    add_table_title(doc, "表6-1  官方AI工具及项目内职责分工")
    add_table(
        doc,
        ["工具名称", "在项目中的角色定位", "具体使用方式"],
        [
            ["腾讯开悟比赛平台", "赛事平台与工具链路入口", "用于赛题资料获取、工具链路打通、作品提交与演示场景对接"],
            ["CodeBuddy", "工程开发与联调工具", "用于前后端页面、接口、状态流与部署脚本的开发协同"],
            ["腾讯元器", "智能体工作流编排主载体", "用于案件识别、法律检索、阶段生成、对方模拟、变量聚合与输出回复"],
            ["腾讯元宝", "Prompt 调优与内容润色辅助工具", "用于优化场景文案、整理表达方式、辅助完善提交材料中的文字呈现"],
            ["小理AI", "法律实务应用参考与表达校验工具", "用于参考法律问答、AI检索与文书类应用形态，辅助法律表达与场景设计优化"],
            ["得理开放平台API", "法律原子能力来源", "用于案例检索、法规检索及法律数据支撑，为元器工作流与后端分析提供专业依据"],
        ],
        col_widths_cm=[2.8, 4.0, 7.8],
    )
    add_body(
        doc,
        "在工具分工上，项目始终坚持“元器负责工作流编排，后端负责业务控制与稳定性，法律能力工具负责专业支撑”的原则。这样做既符合赛题“工作流推荐用元器”的口径，也避免了把元器简单降级成普通接口调用器的误区。同时，通过 CodeBuddy、腾讯元宝、小理AI 等工具参与工程实现和内容优化，项目也形成了一条从开发、调优到材料表达的完整工具链。",
    )

    add_heading(doc, "（二）元器工作流的总体目标", level=2)
    add_body(
        doc,
        "本项目中的腾讯元器工作流，承担的不是单一文本生成，而是对一整条庭审模拟主链路进行编排和约束。它的核心目标包括：第一，识别案件类型与案件中的核心争点；第二，调用法律检索工具获取法规与类案信息；第三，根据当前庭审阶段生成对应内容与对方行为；第四，对多节点结果进行变量聚合，保证前端能够稳定接收结构化输出；第五，为后续的胜诉率分析与复盘报告提供过程数据和结构化上下文。",
    )
    add_body(
        doc,
        "元器在项目中承担流程中枢角色，而不仅是模型文本生成节点。项目借助这一能力，把法律场景中的复杂链路拆解为可被控制的节点，使系统能够在保证 AI 生成能力的同时，尽可能降低流程漂移、变量缺失和输出不稳定的问题。",
    )

    add_heading(doc, "（三）工作流主链路", level=2)
    add_table_title(doc, "表6-2  元器工作流主链路节点设计")
    add_table(
        doc,
        ["阶段顺序", "工作流节点", "主要输入", "主要输出", "作用说明"],
        [
            ["1", "案件识别节点", "案件事实、请求、证据摘要", "案件类型、争议焦点、案件摘要", "识别案件结构，为后续节点确定分析方向"],
            ["2", "法规检索节点", "争点关键词、案件类型", "相关法条与规则结果", "调用得理开放平台API提供法律依据"],
            ["3", "类案检索节点", "案件关键词、争点关键词", "类案裁判要点", "调用得理开放平台API提供类案支撑"],
            ["4", "庭审生成节点", "案件状态、当前阶段、已知争点", "当前阶段场景文本、用户可选动作", "生成当轮庭审互动内容"],
            ["5", "对方行为模拟节点", "当前阶段、对方画像、证据状态", "抗辩意见、补充证据、质证方向", "增强对抗感与不确定性"],
            ["6", "聚合节点", "多节点输出结果", "统一状态对象", "整理空值、统一字段、保证前后端联动稳定"],
            ["7", "结果输出节点", "聚合后的状态对象", "前端展示文本、风险提示、阶段结果", "将结果传给前端与后端进行展示和持久化"],
        ],
        col_widths_cm=[1.1, 2.8, 3.8, 3.8, 4.2],
    )
    add_body(
        doc,
        "在这一主链路中，元器的价值主要体现在三个字：能编排。系统并不是让模型自由发挥，而是让每一个节点各司其职，先识别案件，再取法律依据，再生成当前阶段内容，再模拟对方动作，最后聚合输出结果。这样才能让“庭审模拟”不是一段随机对话，而是一条具备明确输入、明确阶段和明确输出的法律工作流。",
    )

    add_heading(doc, "（四）元器与后端的协同关系", level=2)
    add_body(
        doc,
        "虽然元器承担了智能体工作流编排职责，但本项目并没有把全部逻辑都压到工作流内部处理。原因在于，法律场景中的案件会话、历史记录、阶段推进约束、异常兜底和报告归档，仍然需要由后端进行统一控制。基于此，项目形成了“元器负责智能体与工作流，后端负责状态和稳定性”的协同模式。",
    )
    add_body(
        doc,
        "具体来说，元器负责组织各智能节点之间的关系，输出结构化阶段结果；后端负责维护案件 ID、当前阶段、历史交互记录、证据状态和报告产物，并在元器结果异常、接口延迟或变量不完整时提供兜底逻辑。这样一来，项目既能体现元器作为官方工作流工具的核心地位，又不至于因工作流内变量漂移或接口波动影响整个产品主链路的稳定运行。",
    )

    add_heading(doc, "（五）Prompt设计原则", level=2)
    add_body(
        doc,
        "在 Prompt 设计上，项目坚持“法律专业性优先、结构化输出优先、流程控制优先、风险提示优先”四项原则。第一，所有生成内容都必须围绕真实法律语境展开，不能为了戏剧性牺牲法律逻辑；第二，所有关键节点的输出都应尽量结构化，便于前端展示和后端处理；第三，模型只在当前阶段内生成内容，流程顺序由系统和工作流控制，不交给模型自行决定；第四，对于证据不足、事实不明或结论不确定的情形，Prompt 必须引导模型给出风险提示，而不是输出“保证胜诉”式的不当表述。",
    )
    add_body(
        doc,
        "这套 Prompt 原则在大模型的生成能力和法律场景的专业边界之间建立了明确护栏。Prompt 既要限制泛化回答，也要限制脱离案件事实和程序规则的戏剧化生成。只有明确边界、明确输出结构、明确生成职责，Prompt 才能真正服务于产品。",
    )

    add_heading(doc, "（六）三类核心Prompt设计", level=2)
    add_table_title(doc, "表6-3  核心Prompt设计与控制约束")
    add_table(
        doc,
        ["Prompt类型", "主要职责", "关键输入", "预期输出", "控制约束"],
        [
            ["庭审流程交互 Prompt", "生成当前阶段的庭审叙事与策略动作", "案件摘要、当前阶段、争点、上轮结果", "场景标题、场景描述、用户可选行动、法官提示", "不得跳阶段、不得虚构核心事实、输出需结构化"],
            ["对方行为模拟 Prompt", "模拟对方律师/当事人/证人的抗辩与质证路径", "对方画像、证据状态、当前争点", "抗辩意见、突袭证据、程序动作", "必须与案件设定一致，不能脱离法律逻辑制造冲突"],
            ["胜诉率分析 Prompt", "基于全流程结果输出胜算评估与准备建议", "争点、证据、法条、类案、模拟过程", "分维度分析、风险点、证据缺口、行动建议", "不得给出绝对承诺，必须明确分析依据与不确定性"],
        ],
        col_widths_cm=[3.3, 3.4, 3.3, 3.6, 3.0],
    )

    add_heading(doc, "（七）Prompt示意内容", level=2)
    add_body(
        doc,
        "项目将核心 Prompt 的呈现方式控制在“职责说明 + 输入范围 + 输出要求 + 约束条件”四个维度，而不是直接贴出大段原始提示词。以庭审流程交互 Prompt 为例：系统首先告知模型当前扮演的角色及所处庭审阶段，然后注入结构化案件信息和争点摘要，再明确要求模型生成当前阶段内的法庭文本、用户可选行动和一句法官提示，最后强调禁止跳阶段、禁止虚构关键事实、禁止输出绝对结论。这样的设计既保留了法律表达的灵活性，也保证了系统输出可控。",
    )
    add_body(
        doc,
        "对方行为模拟 Prompt 则重点约束“只在现有案件事实与证据边界内生成对抗内容”，使模型更多体现真实法庭中的质疑、抗辩与程序动作，而不是脱离案件设定进行任意扩写。胜诉率分析 Prompt 则更强调“解释性”，即要求模型必须分维度说明何以支持、何以不足、哪些地方存在风险、哪些证据需要补强，从而使结果不是一句抽象判断，而是一组具有实际指导意义的分析输出。",
    )

    add_heading(doc, "（八）风险控制与人工审校机制", level=2)
    add_body(
        doc,
        "为了保证 AI 输出不偏离法律应用的基本要求，项目在 AI 工具使用上设置了双层控制机制。第一层是系统控制，即通过元器工作流、后端状态机、变量聚合和结构化输出要求，减少大模型任意发挥导致的失真风险；第二层是人工审校，即由法学同学对法律术语、程序表述、法条引用和模拟口径进行人工审核，确保项目中的关键法律内容不会因生成式表达而偏离基本专业边界。",
    )
    add_body(
        doc,
        "在这一机制下，AI 的定位始终是“辅助分析与交互生成工具”，而不是对真实案件结果作出保证性判断的主体。项目在所有相关说明中都强调：系统输出仅作为庭前准备、学习训练和模拟推演的辅助结果，不能替代正式法律意见。也正因为具备这样的边界意识和控制机制，项目才能在体现 AI 创新能力的同时，保持智慧法律应用所要求的专业稳健性。",
    )


def build_section_value_outlook(doc: Document) -> None:
    add_heading(doc, "第七章  项目价值、创新点与实施展望", level=1)
    add_body(
        doc,
        "从项目价值层面看，“企鹅法庭”最大的意义不在于做出一个会说法律术语的智能体，而在于把法律场景中最难被提前练习的一段过程，即庭前准备与庭审对抗预判，转化为一个可以被普通用户理解、被法学生重复训练、被评委直观看到成果的交互式系统。它所回应的是法律服务中的真实问题，因此具有较强的应用解释力和展示价值。",
    )
    add_heading(doc, "（一）项目综合价值", level=2)
    add_body(
        doc,
        "在社会价值层面，该系统有助于降低普通人进入诉讼准备阶段的认知门槛，使用户能够更清晰地理解自己的案件结构和证据状态；在教学价值层面，该系统能够作为模拟法庭和法律实训课程的辅助工具，为学生提供低成本、可重复、可对比的训练环境；在赛事价值层面，该系统提供了一种较为完整的智慧法律系统解决方案表达路径，能够更好体现赛题对于“法律需求—技术方案—工作流设计—演示闭环”的要求。",
    )
    add_heading(doc, "（二）创新点再归纳", level=2)
    add_body(
        doc,
        "综合全文，项目创新点可再归纳为三方面：其一，交互组织方式的创新，即以阶段化法庭过程替代普通对话窗口，提升沉浸感与任务导向性；其二，系统链路设计的创新，即将案件录入、对方模拟、胜算分析和结果报告纳入一条受控主轴；其三，赛事工具链落地方式的创新，即以腾讯元器为核心进行工作流编排，同时通过后端控制层保障系统运行稳定，使作品既满足赛题工具要求，又保留工程可交付性。",
    )
    add_heading(doc, "（三）当前边界与后续展望", level=2)
    add_body(
        doc,
        "当前版本仍以区域初赛可交付、可演示和主链路可跑通为首要目标，因此在案件类型覆盖面、复杂程序适配、多人协同模式和更高阶的法律解释能力方面，仍保留后续扩展空间。下一阶段可沿着两个方向继续完善：一是补强案件类型、证据规则模板和标准化输入输出样例，提高系统专业覆盖度；二是优化工作流联动、结果追踪与演示部署，使系统在更稳定的工程形态上走向正式展示与更长期的科创项目沉淀。",
    )
    add_heading(doc, "（四）结论", level=2)
    add_body(
        doc,
        "综上所述，“企鹅法庭·沉浸式庭审模拟与法律辅助分析系统”并非简单的法律问答工具，而是一套围绕诉前准备与庭审推演构建的智慧法律系统解决方案。项目在赛题工具链约束下，完成了从需求分析、工作流设计、核心功能规划到应用场景落地的系统化表达，已经形成较清晰的问题意识、技术路径和成果闭环，可用于区域初赛阶段的正式申报与后续答辩表达。",
    )


def build_references(doc: Document) -> None:
    add_heading(doc, "参考文献", level=1)
    references = [
        "腾讯开悟. 赛道详情：法律AI应用创新与实践[EB/OL]. https://tencentarena.com/aiarena/zh/match/open-competition-2026/open-competition-2026-6, 2026-04-20.",
        "腾讯开悟. 腾讯开悟 - Tencent AI Arena[EB/OL]. https://tencentarena.com/aiarena/zh, 2026-04-20.",
        "腾讯元器. 腾讯元器介绍[EB/OL]. https://yuanqi.tencent.com/guide/yuanqi-introduction, 2026-04-20.",
        "腾讯元器. API接口文档[EB/OL]. https://yuanqi.tencent.com/guide/publish-agent-api-documentation, 2026-04-20.",
        "全国人民代表大会常务委员会. 全国人民代表大会常务委员会关于修改《中华人民共和国民事诉讼法》的决定[EB/OL]. https://www.npc.gov.cn/npc/c2/c30834/202309/t20230901_431419.html, 2023-09-01.",
        "全国人民代表大会常务委员会. 全国人民代表大会常务委员会关于修改《中华人民共和国民事诉讼法》的决定[EB/OL]. https://www.npc.gov.cn/c2/c30834/202112/t20211224_315578.html, 2021-12-24.",
        "深圳得理科技有限公司. 得理 - 法律行业智能化变革引领者[EB/OL]. https://www.delilegal.com/, 2026-04-20.",
        "深圳得理科技有限公司. 得理开放平台[EB/OL]. https://open.delilegal.com/home/start, 2026-04-20.",
        "企鹅法庭项目组. 企鹅法庭项目技术方案V1[R]. 2026.",
        "企鹅法庭项目组. AI工具使用说明（企鹅法庭）[R]. 2026.",
    ]
    for idx, item in enumerate(references, start=1):
        add_reference_item(doc, idx, item)


def build_appendix(doc: Document) -> None:
    add_heading(doc, "附录  区域初赛材料映射表", level=1)
    add_body(
        doc,
        "本附录列示区域初赛“项目简介”要求与本报告具体章节的对应关系，可直接定位各评审关注点在文稿中的呈现位置，便于后续拆分材料、制作PPT和准备答辩。",
    )
    add_table_title(doc, "表A-1  区域初赛项目简介要求与本报告章节映射")
    add_table(
        doc,
        ["赛题要求项", "本报告对应章节", "说明"],
        [
            ["项目基础信息", "第二章 项目基础信息与总体定位", "包含项目名称、定位、目标对象与总体架构"],
            ["项目需求分析", "第三章 项目需求分析与问题建模", "包含用户画像、痛点拆解与优先级判断"],
            ["关键功能说明", "第四章 系统设计与关键功能实现", "包含案件录入、庭审模拟、分析报告等核心模块"],
            ["项目使用场景", "第五章 项目使用场景与演示验证", "包含普通用户、法学生、普法展示和比赛答辩场景"],
            ["AI工具使用说明（工作流与Prompt）", "第六章 AI工具链、工作流与Prompt设计", "包含工具分工、工作流主链路、Prompt原则与风险控制"],
        ],
        col_widths_cm=[4.0, 6.0, 5.4],
    )


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    build_cover(doc)

    front_matter = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section_layout(front_matter)
    set_page_number_start(front_matter, 1)
    set_section_header(front_matter, None)
    set_section_footer(front_matter, roman=True, visible=True)
    build_abstract(doc)
    build_catalog(doc)

    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section_layout(main_section)
    set_page_number_start(main_section, 1)
    set_section_header(main_section, "企鹅法庭·沉浸式庭审模拟与法律辅助分析系统")
    set_section_footer(main_section, roman=False, visible=True)
    build_introduction(doc)
    build_section_basic_info(doc)
    build_section_needs(doc)
    build_section_features(doc)
    build_section_scenarios(doc)
    build_section_ai_tools(doc)
    build_section_value_outlook(doc)
    build_references(doc)
    build_appendix(doc)
    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
