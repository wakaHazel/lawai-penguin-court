import sys
from docx import Document

def read_docx(file_path):
    doc = Document(file_path)
    print(f"Reading: {file_path}")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    for i, p in enumerate(doc.paragraphs[:20]):  # Print first 20 paragraphs to inspect structure
        if p.text.strip():
            print(f"P{i} [{p.style.name}]: {p.text}")
    print("...")

if __name__ == "__main__":
    read_docx("E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_论文标准版_修订版.docx")
    print("\n===============================\n")
    read_docx("E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_论文标准版_修订版.docx")
