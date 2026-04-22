import os
import docx

def deep_clean_presentation(file_path, out_path):
    print(f"--- 深度清洗 {os.path.basename(file_path)} ---")
    doc = docx.Document(file_path)
    
    replacements = {
        # 项目简介中的遗留
        "演示链路": "核心业务链路",
        "项目使用场景与演示验证": "项目应用场景与商业验证",
        "法律援助、普法与体验展示场景": "法律援助与普法服务场景",
        
        # Demo文档中的大修
        "项目Demo与补充材料研究文稿": "项目功能验证与技术白皮书",
        "项目演示方案与补充材料结构说明": "核心功能验证与技术实现结构说明",
        "演示 方案": "功能验证方案",
        "演示方案": "功能验证方案",
        "Demo 的展示逻辑": "MVP版本的业务逻辑",
        "Demo 与补充材料": "功能原型与技术白皮书",
        "Demo 在整体产品验证中的定位": "功能原型在整体产品验证中的定位",
        "展示系统界面": "呈现系统界面",
        "Demo展示交互": "原型交互体验",
        "提交材料": "技术文档",
        "项目 Demo": "产品 MVP (最小可行性产品)",
        "演示链路中的关键环节": "业务链路中的关键环节",
        "Demo  是评估": "MVP 是评估",
        "Demo 是评估": "MVP 是评估",
        "Demo 演示主线": "核心功能主线",
        "演示过程": "验证过程",
        "展示庭审文本": "生成庭审文本",
        "展示完整的系统闭环": "形成完整的系统闭环",
        "演示地址": "体验地址",
        "采用本地演示": "采用本地化部署",
        "本地稳定演示 版本": "本地稳定运行版本",
        "本地稳定演示版本": "本地稳定运行版本",
        "视频演 示与界面展示": "运行录屏与界面说明",
        "视频演示与界面展示": "运行录屏与界面说明",
        "Demo核心演示内容": "核心功能验证内容",
        "Demo 的演示路径": "MVP 的验证路径",
        "集中展示系统": "集中体现系统",
        "主线演示": "主线流程",
        "演示视频": "产品操作视频",
        "Demo演示案例": "标准验证案例",
        "演示样本": "验证样本",
        "演示案例": "验证案例",
        "重点演示": "重点验证",
        "Demo提交材料清单": "技术白皮书附录清单",
        "可信度的展示": "可信度的说明",
        "简要展示": "简要说明",
        "展示系统结果": "体现系统结果",
        "演示稿": "概念稿",
        "一次性演示稿": "一次性概念稿",
        "本地演示版本稳定可运行": "本地部署版本稳定可运行"
    }

    for p in doc.paragraphs:
        original_text = p.text
        new_text = original_text
        for old_str, new_str in replacements.items():
            if old_str in new_text:
                new_text = new_text.replace(old_str, new_str)
        
        if original_text != new_text:
            print(f"修改前: {original_text}")
            print(f"修改后: {new_text}\n")
            p.clear()
            p.add_run(new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original_text = p.text
                    new_text = original_text
                    for old_str, new_str in replacements.items():
                        if old_str in new_text:
                            new_text = new_text.replace(old_str, new_str)
                    if original_text != new_text:
                        p.clear()
                        p.add_run(new_text)

    doc.save(out_path)
    print(f"✅ 保存到: {out_path}\n")

files = [
    r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_去比赛痕迹_纯净版.docx",
    r"E:\lawai\output\doc\区域初赛提交材料_项目Demo与补充材料_企鹅法庭_去比赛痕迹_纯净版.docx"
]

for f in files:
    out_path = f.replace("去比赛痕迹_纯净版", "商业白皮书_最终版")
    deep_clean_presentation(f, out_path)
