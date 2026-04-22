import os
import docx

def remove_tables_add_placeholder(file_path, out_path, prefix):
    print(f"--- 替换 {os.path.basename(file_path)} 的表格 ---")
    doc = docx.Document(file_path)
    
    for i, table in enumerate(doc.tables):
        # 插入一段占位符提示，代替表格
        p = table._element.getprevious()
        if p is not None and p.tag.endswith('p'):
            para = docx.text.paragraph.Paragraph(p, doc._body)
            # 在表格前面插入一个高亮的提示文本，告诉用户在这里拖入哪张 SVG
            svg_filename = f"{prefix}_table_{i+1}.svg"
            
            run = para.add_run(f"\n[📝 请在此处插入精美图表： diagrams/tables/{svg_filename} ]")
            run.font.color.rgb = docx.shared.RGBColor(0xFF, 0x00, 0x00) # 红色高亮
            run.bold = True
            
        # 物理删除原表格的 XML 节点
        table._element.getparent().remove(table._element)

    doc.save(out_path)
    print(f"✅ 保存到: {out_path}")

files_to_process = {
    r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_科创展示版_最终定稿.docx": "intro",
    r"E:\lawai\output\doc\区域初赛提交材料_项目Demo与补充材料_企鹅法庭_科创展示版_最终定稿.docx": "demo"
}

for f_path, prefix in files_to_process.items():
    out_path = f_path.replace("最终定稿", "无表格_占位版")
    remove_tables_add_placeholder(f_path, out_path, prefix)
