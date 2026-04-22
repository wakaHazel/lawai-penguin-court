import os
import docx

def replace_in_text(text):
    replacements = {
        "在比赛层面，系统要满足赛题对腾讯系工具使用、工作流与 Prompt 说明、专业验证测试和完整演示闭环的要求，形成一套可展示、可答辩、可复现的智慧法律应用方案。": "形成一套可展示、可验证、可复现的智慧法律应用方案。",
        "这条主链路打通后，项目即具备产品展示和赛题答辩的核心价值。": "这条主链路打通后，项目即具备产品展示与落地的核心价值。",
        "与项目简介、视频封面和答辩 PPT 保持同名": "与项目简介、视频封面保持同名",
        "作为视频和现场答辩固定主线": "作为视频和现场演示固定主线",
        "便于后续彩排与答辩统一口径": "便于后续功能演示统一口径",
        "答辩问答库或演示脚本": "常见问题库与演示脚本",
        "形成 完整主链路、界面交互和报告成果，便于展示与答辩": "形成完整主链路、界面交互和报告成果，便于展示与验证",
        "比赛 答辩": "产品 展示",
        "包含 普通用户、法学生、普法展示和比赛答辩场景": "包含普通用户、法学生和普法展示场景",
        "AI工 具使用说明": "大模型工具与技术说明",
        "AI工具使用说明": "大模型工具与技术说明",
        "答辩": "演示",
        "评委": "专家",
        "赛题": "项目",
        "比赛": "科创",
    }
    
    new_text = text
    # Exact match replacements first
    for old, new in replacements.items():
        if old in new_text:
            new_text = new_text.replace(old, new)
            
    return new_text

def clean_doc(file_path):
    print(f"Cleaning: {file_path}")
    doc = docx.Document(file_path)
    
    # 1. Clean paragraphs and delete fake references
    paragraphs_to_delete = []
    for p in doc.paragraphs:
        # Delete fake references completely
        if '企鹅法庭项目技术方案' in p.text or 'AI工具使用说明' in p.text and '[09]' in p.text or '[10]' in p.text:
            paragraphs_to_delete.append(p)
            continue
            
        # Replace text in paragraphs
        if '答辩' in p.text or '评委' in p.text or '比赛' in p.text or '赛题' in p.text:
            # paragraph level rewrite (loses some formatting but ensures text is replaced across run boundaries)
            full_text = p.text
            new_text = replace_in_text(full_text)
            for r in p.runs:
                r.text = ""
            if len(p.runs) > 0:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)

    # Delete marked paragraphs
    for p in paragraphs_to_delete:
        p._element.getparent().remove(p._element)

    # 2. Clean tables
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if '答辩' in p.text or '评委' in p.text or 'AI工具' in p.text or 'AI工 具' in p.text or '比赛' in p.text or '赛题' in p.text:
                        full_text = p.text
                        new_text = replace_in_text(full_text)
                        for run in p.runs:
                            run.text = ""
                        if len(p.runs) > 0:
                            p.runs[0].text = new_text
                        else:
                            p.add_run(new_text)

    # Save to a new file to avoid permission error if user has it open
    new_file_path = file_path.replace('.docx', '_Cleaned.docx')
    doc.save(new_file_path)
    print(f"Saved: {new_file_path}")

docs = [
    'E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_科创展示版_最终定稿.docx',
    'E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_科创展示版_最终定稿.docx'
]

for d in docs:
    if os.path.exists(d):
        clean_doc(d)
