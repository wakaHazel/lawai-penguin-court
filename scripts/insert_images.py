import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_image_after_text(doc, search_text, image_path, caption_text):
    for idx, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            # Check if image is already inserted in the next paragraph
            if idx + 1 < len(doc.paragraphs) and caption_text in doc.paragraphs[idx+1].text:
                print(f"Image already exists for: {caption_text}")
                return
            if idx + 2 < len(doc.paragraphs) and caption_text in doc.paragraphs[idx+2].text:
                print(f"Image already exists for: {caption_text}")
                return
                
            # Create a new paragraph for the image
            # We must use oxml to insert the paragraph directly after the current one
            new_p = OxmlElement('w:p')
            p._p.addnext(new_p)
            img_p = doc.paragraphs[idx + 1]
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = img_p.add_run()
            try:
                run.add_picture(image_path, width=Inches(5.5))
            except Exception as e:
                print(f"Failed to insert image {image_path}: {e}")
                return
                
            # Create a new paragraph for the caption
            new_caption_p = OxmlElement('w:p')
            img_p._p.addnext(new_caption_p)
            caption_p = doc.paragraphs[idx + 2]
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            caption_run = caption_p.add_run(caption_text)
            caption_run.font.size = Pt(10)
            caption_run.font.bold = True
            
            # Set spacing
            caption_p.paragraph_format.space_after = Pt(12)
            
            print(f"✅ Inserted image {image_path} with caption '{caption_text}'")
            return
            
    print(f"⚠️ Could not find text '{search_text}' in document.")

def main():
    base_cg_dir = "E:/lawai/data/cg-library/cartoon-court"
    img_prepare = os.path.join(base_cg_dir, "stage_prepare.png")
    img_debate = os.path.join(base_cg_dir, "stage_debate.png")
    img_report = os.path.join(base_cg_dir, "stage_report_ready.png")
    
    # 1. 插入到项目简介中
    intro_doc_path = "E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_流畅严谨版.docx"
    intro_out_path = "E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_图文并茂版.docx"
    
    if os.path.exists(intro_doc_path):
        doc = Document(intro_doc_path)
        insert_image_after_text(doc, "系统以民事诉讼庭前准备阶段为核心，建立了“案件录入—庭审推演—对手模拟—胜率评估—复盘报告”的完整工作流", img_prepare, "图 1：企鹅法庭·诉前准备与案件推演工作台界面模拟")
        insert_image_after_text(doc, "在应用设计上，系统将常规的诉讼流程转化为分阶段、多分支的互动形式。用户可以在高度还原的法庭场景中", img_debate, "图 2：企鹅法庭·沉浸式庭审对抗与举证质证推演")
        insert_image_after_text(doc, "从目前的系统实现与演示链路来看，该项目在诉前辅助准备、法学实践教学等场景中具有良好的适用性", img_report, "图 3：企鹅法庭·胜诉率分析与全景复盘报告系统")
        doc.save(intro_out_path)
        print(f"Saved: {intro_out_path}")
        
    # 2. 插入到Demo与补充材料中
    demo_doc_path = "E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_流畅严谨版.docx"
    demo_out_path = "E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_图文并茂版.docx"
    
    if os.path.exists(demo_doc_path):
        doc = Document(demo_doc_path)
        insert_image_after_text(doc, "目前，“企鹅法庭”已开发为完整的 Web 端应用系统。其业务主线遵循“案件录入—庭审推演—互动质证—胜率分析—复盘输出”的逻辑顺序。", img_prepare, "图 1：系统核心链路演示——案件录入与初始化推演阶段")
        insert_image_after_text(doc, "演示过程采用“预置案例 + 实时互动 + 结果输出”的模式。首先通过导入预设的案件模板快速进入工作台；随后通过一至两轮的核心节点交互", img_debate, "图 2：典型案例演示——多模态庭审质证与对手策略响应")
        doc.save(demo_out_path)
        print(f"Saved: {demo_out_path}")

if __name__ == "__main__":
    main()
