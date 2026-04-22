import os
import docx
import difflib

def extract_text_simple(path: str):
    if not os.path.exists(path):
        print(f"❌ 找不到文件: {path}")
        return []
    doc = docx.Document(path)
    # 只提取非空的段落和表格内容，方便对比
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in lines:  # 简单去重防止重复提取
                    lines.append(text)
    return lines

def compare_documents():
    file_original = r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_流畅严谨版.docx"
    file_modified = r"C:\Users\温岐\Desktop\区域初赛提交材料_项目简介_企鹅法庭（1改）.docx"
    
    print(f"--- 正在对比文档差异 ---")
    print(f"原始版本: {os.path.basename(file_original)}")
    print(f"修改版本: {os.path.basename(file_modified)}\n")
    
    text1 = extract_text_simple(file_original)
    text2 = extract_text_simple(file_modified)
    
    if not text1 or not text2:
        return
        
    diff = list(difflib.unified_diff(text1, text2, n=0))
    
    if not diff:
        print("✅ 两份文档内容完全一致，没有实质性文本修改。")
        return
        
    print(f"发现差异，具体修改如下：\n")
    
    changes_count = 0
    for line in diff:
        if line.startswith('---') or line.startswith('+++') or line.startswith('@@'):
            continue
        elif line.startswith('-'):
            print(f"🔴 删除了: {line[1:]}")
            changes_count += 1
        elif line.startswith('+'):
            print(f"🟢 新增/修改为: {line[1:]}\n")
            
    print(f"\n对比完成。共发现约 {changes_count} 处差异。")

if __name__ == "__main__":
    compare_documents()
