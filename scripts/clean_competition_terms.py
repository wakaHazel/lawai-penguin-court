import os
import docx

def clean_document(file_path, out_path):
    print(f"--- 清洗 {os.path.basename(file_path)} ---")
    doc = docx.Document(file_path)
    
    # 精确替换字典（长句优先）
    replacements = {
        "区域初赛项目研究报告": "项目研究报告",
        "提交阶段：区域初赛": "当前阶段：V1.0版发布",
        "赛题《法律AI应用创新与实践》明确鼓励参赛团队": "本项目依托前沿AI工具平台，明确鼓励研发团队",
        "在比赛层面": "在工程验证层面",
        "满足赛题对腾讯系工具使用": "满足对前沿工具使用",
        "可展示、可答辩、可复现": "可展示、可落地、可复现",
        "产品展示和赛题答辩": "产品展示和商业落地",
        "比赛周期": "产品研发周期",
        "对于区域初赛来说，评委首先关心的是": "对于一款成熟的法律科技产品而言，核心在于",
        "被评委直观看到": "被行业专家直观看到",
        "区域初赛可交付": "MVP（最小可行性产品）可交付",
        "区域初赛阶段的正式申报与后续答辩表达": "后续的商业化落地与规模化应用奠定了坚实基础",
        "区域初赛Demo与补充材料研究文稿": "项目Demo与补充材料研究文稿",
        "项目演示方案、补充材料结构与答辩支撑内容": "项目演示方案与补充材料结构说明",
        "作为区域初赛的配套说明": "作为本项目的配套说明",
        "支撑评委对项目": "支撑受众对项目",
        "Demo 在评审中的定位": "Demo 在整体产品验证中的定位",
        "项目在评审环节形成": "项目在验收环节形成",
        "答辩支撑；提交设计": "系统设计；应用落地",
        "在初赛评审中": "在产品验收中",
        "若比赛阶段采用": "若当前阶段采用",
        "初赛的演示案例": "初版MVP的演示案例",
        "对法律 AI 赛题而言 ，评审通常重点关注": "对法律 AI 产品而言，行业专家通常重点关注",
        "对法律 AI 赛题而言，评审通常重点关注": "对法律 AI 产品而言，行业专家通常重点关注",
        "有助于评委确认": "有助于专业受众确认",
        "对应赛题对工具": "对应了行业对技术栈",
        "等评审关注问题": "等专业关注问题",
        "使评委在不亲自": "使用户在不亲自",
        "和答辩准备": "和产品发布准备",
        "区域初赛材料映射表": "项目材料架构表",
        "区域初赛项目简介要求与本报告章节映射": "项目要求与本报告章节映射",
        "便于后续拆分材料、制作PPT和准备答辩": "便于后续拆分材料和制作产品演示PPT",
        # 兜底替换
        "区域初赛": "当前阶段",
        "初赛": "初期",
        "参赛团队": "研发团队",
        "赛题": "项目需求",
        "答辩": "成果展示",
        "评委": "专家",
        "评审": "验收",
        "比赛": "研发"
    }

    # 找到附录所在的段落并删除（如果存在）
    # 由于 python-docx 删除段落稍微有点麻烦，我们采用清空内容的方式
    deleting_appendix = False
    for p in doc.paragraphs:
        if "附录" in p.text and "映射表" in p.text:
            deleting_appendix = True
        
        if deleting_appendix:
            # 删除段落的底层 XML 节点
            p._element.getparent().remove(p._element)
            continue

        # 文本替换（保留原有样式，只替换文本）
        original_text = p.text
        new_text = original_text
        for old_str, new_str in replacements.items():
            if old_str in new_text:
                new_text = new_text.replace(old_str, new_str)
        
        if original_text != new_text:
            print(f"修改前: {original_text}")
            print(f"修改后: {new_text}\n")
            # 保留段落格式的替换
            p.clear()
            p.add_run(new_text)

    # 处理表格中的文本
    for table in doc.tables:
        if deleting_appendix: # 如果整个附录被删，可能表格也需要删掉
            table._element.getparent().remove(table._element)
            continue
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original_text = p.text
                    new_text = original_text
                    for old_str, new_str in replacements.items():
                        if old_str in new_text:
                            new_text = new_text.replace(old_str, new_str)
                    if original_text != new_text:
                        p.clear()
                        p.add_run(new_text)

    doc.save(out_path)
    print(f"✅ 保存到: {out_path}\n")

files = [
    r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_架构图文版.docx",
    r"E:\lawai\output\doc\区域初赛提交材料_项目Demo与补充材料_企鹅法庭_架构图文版.docx"
]

for f in files:
    out_path = f.replace("架构图文版", "去比赛痕迹_纯净版")
    clean_document(f, out_path)
