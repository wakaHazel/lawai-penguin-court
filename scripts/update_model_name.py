import os
import docx

def replace_model_name(file_path):
    print(f"--- 更新 {os.path.basename(file_path)} ---")
    doc = docx.Document(file_path)
    
    replacements = {
        "Gemini-3.1-Pro": "CodeBuddy LLM",
        "Gemini LLM": "CodeBuddy LLM",
        "Gemini": "CodeBuddy"
    }
    
    changed = False

    for p in doc.paragraphs:
        original_text = p.text
        new_text = original_text
        for old_str, new_str in replacements.items():
            if old_str in new_text:
                new_text = new_text.replace(old_str, new_str)
        
        if original_text != new_text:
            print(f"修改前: {original_text}")
            print(f"修改后: {new_text}\n")
            p.clear()
            p.add_run(new_text)
            changed = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original_text = p.text
                    new_text = original_text
                    for old_str, new_str in replacements.items():
                        if old_str in new_text:
                            new_text = new_text.replace(old_str, new_str)
                    if original_text != new_text:
                        p.clear()
                        p.add_run(new_text)
                        changed = True

    if changed:
        doc.save(file_path)
        print(f"✅ 已保存修改: {file_path}\n")
    else:
        print("未发现需要替换的文本\n")

files = [
    r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_科创展示版_最终定稿.docx",
    r"E:\lawai\output\doc\区域初赛提交材料_项目Demo与补充材料_企鹅法庭_科创展示版_最终定稿.docx"
]

for f in files:
    if os.path.exists(f):
        replace_model_name(f)
    else:
        print(f"找不到文件: {f}")
