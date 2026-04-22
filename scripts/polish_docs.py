import os
import json
import logging
from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Since LLM calls inside python script are hard without a client configured, 
# I will use a prompt-based transformation logic, but simulating a much more natural, 
# engaging and less "engineering-stiff" tone. I will write a simple replacement 
# dictionary for common stiff phrases, and also restructure some sentences.

REPLACEMENTS = {
    "随着法律服务数字化进程的持续推进，法律行业对于能够兼顾专业性、可解释性与可交互性的智慧法律应用提出了更高要求。": 
        "在法律服务日益数字化的今天，行业不仅需要冰冷的法条检索，更呼唤那些兼具专业深度、透明度与人性化交互的智慧法律助手。",
        
    "针对普通诉讼参与人难以在庭前有效整理案情、预判对方抗辩和评估证据强弱的问题":
        "很多普通人在打官司前，往往面对一团乱麻的案情不知从何下手，更难以预判对方的招数或评估自己手里的证据。为了打破这种困境",
        
    "项目围绕【D06】法律AI应用创新与实践赛 题构建了“企鹅法庭·沉浸式庭审模拟与法律辅助分析系统”。":
        "我们紧扣【D06】法律AI应用创新与实践赛题，倾力打造了“企鹅法庭”——一款沉浸式的庭审模拟与法律辅助分析系统。",
        
    "该系统以民事诉讼庭前准备场 景为核心，构建了“案件录入—庭审模拟—对方行为推演—胜诉率分析—复盘报告”的完整主轴，在产品形态上突破了传统法律问答工具停留于单点回答的局限。":
        "这套系统聚焦于民事诉讼的庭前准备阶段。我们精心设计了“案件录入—庭审推演—对手模拟—胜率评估—复盘总结”的闭环体验，让产品不再是一个只会单句问答的机器人，而是一位真正能陪你全程演练的“数字律师”。",
        
    "在技术方案上，项目采用“前端交互层—后端控制层—腾讯元器工作流编排 层—法律能力工具层”的组合架构":
        "在底层架构上，系统优雅地融合了四层设计：灵动的前端交互、稳健的后端中枢、腾讯元器驱动的智能工作流，以及专业的法律知识引擎。",
        
    "以腾讯元器承担流程编排与智能体协同，以得理开放平台API和小理AI承担法律检索与专业支撑，以后端状态控制保证多阶段庭审模拟的稳定推进。":
        "其中，腾讯元器像一位经验丰富的导演，精准调度着各个智能体的协作；得理开放平台与小理AI则作为强大的专业后盾，提供精准的法律检索支撑；而我们自研的后端状态机，则确保了这场多阶段的“模拟法庭”能够有条不紊地推进。",
        
    "在应用设计上，系统将民事诉讼典型流程转化为分阶段、可分支、可复盘的文游式互动过程，使用户能够在接近真实法庭语境的环境中进行策略选择、对抗预判和准备校正。":
        "在体验设计上，我们别出心裁地将枯燥的诉讼流程，转化为了类似“文字冒险游戏（文游）”的互动体验。用户可以在高度还原的法庭氛围中，像玩游戏一样步步为营：选择应对策略、直面对手质证、并随时调整自己的庭审准备。",
        
    "从当前系统实现、工作流设计与演示链路看，项目在普通当事人诉前准备、法学生模拟法庭训练以及智慧法律展示等场景中已经形成较强的适配性和展示潜力，也为法律AI系统从“问答型”向“ 过程型、策略型、结果沉淀型”演进提供了可落地的系统样本。":
        "无论是帮普通老百姓做诉前演练，还是给法学生提供逼真的模拟法庭训练，亦或是作为智慧法律的创新展示，“企鹅法庭”都展现出了极大的落地价值。它向大家证明：法律AI完全可以从简单的“一问一答”，进化为“重过程、讲策略、出成果”的智能伙伴。",
        
    "作为区域初赛配套文稿，本研究文稿围绕“企鹅法庭·沉浸式庭审模拟与法律辅助分析系统”的项目演示方案与补充材料组织方式展开，重点说明作品 Demo 的展示主线 、视频脚本、案例选取原则、提交项构成，以及补充材料如何服务于项目完成度、专业性和答辩说服力的表达。":
        "本配套文稿是“企鹅法庭”项目的全景导览。我们将为您详细拆解作品 Demo 的展示脉络、视频设计与案例挑选逻辑，并向评委清晰展示：这些精心准备的补充材料，是如何全方位印证项目的完成度、专业底蕴以及实战价值的。",
        
    "文稿遵循“展示主轴清晰、证据材料可核、系统能力可映射、评审问题 可直接响应”的组织原则，将项目 Demo 与补充材料从单纯的提交附件提升为支撑作品理解 、作品验证与现场答辩的正式配套研究材料。":
        "我们在准备这些材料时，始终坚持“主线清晰、查证有据、能力可见、直击痛点”的原则。对我们而言，Demo与补充材料绝不是走过场的附件，而是帮助大家深度理解作品、验证技术实力、以及支撑现场答辩的核心利器。",
        
    "文稿首先分析作品 Demo 在区域初赛中的功能定位，明确其不仅承担展示系统界面的作用，更承担验证系统完成度、解释主工作流、证明工具链落地与引导演示节奏的任务；":
        "首先，作品 Demo 不仅仅是用来“看界面”的。它是整个系统完成度的试金石，更是我们向大家直观解释核心工作流、证明腾讯AI工具链切实落地、并把控演示节奏的关键媒介。",
        
    "继而从项目技术方案、AI 工具与工作流说明、输入输出示例包、团队分工与专业验 证说明、部署与运行说明等方面，构建补充材料的总体结构。":
        "紧接着，我们通过详实的技术方案、清晰的AI工作流图纸、丰富的输入输出用例，以及明确的团队分工与部署指南，构建了一个立体、透明的补充材料体系。",
        
    "通过上述设计，项目可以在正式提交与答辩场景中形成“主报告负责问题与方案、Demo负责系统展示、补充材料负责证据 支撑”的三层表达体系，从而提高区域初赛材料的整体完整性与学术规范性。":
        "有了这样精心的设计，我们的答辩将形成一套无懈可击的“三板斧”：主报告讲透痛点与方案，Demo惊艳展示系统实力，补充材料提供硬核的数据支撑。这不仅极大地提升了项目的完整度，也充分展现了我们严谨的参赛态度。"
}

def polish_text(text: str) -> str:
    if not text.strip():
        return text
        
    polished = text
    for old_str, new_str in REPLACEMENTS.items():
        # Remove extra spaces in keys for robust matching, or just do simple replace
        # Since docx might have weird line breaks, let's clean the text for matching
        clean_old = old_str.replace(" ", "")
        clean_text = polished.replace(" ", "")
        
        if clean_old in clean_text:
            # We found a match, but we need to replace it in the original text while preserving original spacing as much as possible?
            # Easier way: just do a direct string replacement if it's an exact match.
            pass
            
        # Fallback to direct replace
        polished = polished.replace(old_str, new_str)
        
        # Also handle cases where docx has spaces inside words due to justification
        old_with_space = old_str.replace("赛 题", "赛题").replace("场 景", "场景").replace("编排 层", "编排层").replace("问题 可", "问题可").replace("证据 支撑", "证据支撑")
        polished = polished.replace(old_with_space, new_str)
        
    return polished

def process_docx(input_path: str, output_path: str):
    doc = Document(input_path)
    
    # Process paragraphs
    for p in doc.paragraphs:
        # We need to replace text at the run level to preserve formatting, 
        # but text replacement across runs is notoriously difficult in python-docx.
        # A simpler approach is to replace the text of the entire paragraph 
        # and re-apply the style, OR if the paragraph is mostly uniform, just set p.text.
        # Let's try to set p.text but preserve style.
        
        original_text = p.text
        if not original_text.strip():
            continue
            
        # Clean up weird spaces in Chinese text that docx sometimes adds
        clean_text = original_text.replace("赛 题", "赛题").replace("场 景", "场景").replace("编排 层", "编排层").replace("主线 、", "主线、").replace("问题 可", "问题可").replace("证据 支撑", "证据支撑").replace("专业验 证", "专业验证")
            
        polished_text = clean_text
        for old_str, new_str in REPLACEMENTS.items():
            # Strip spaces from old_str to match the cleaned text
            clean_old = old_str.replace("赛 题", "赛题").replace("场 景", "场景").replace("编排 层", "编排层").replace("问题 可", "问题可").replace("证据 支撑", "证据支撑")
            polished_text = polished_text.replace(clean_old, new_str)
            
        if polished_text != original_text:
            # Save paragraph formatting
            style = p.style
            alignment = p.alignment
            
            # Clear runs and add new text
            p.clear()
            new_run = p.add_run(polished_text)
            
            # Note: We might lose bold/italic on specific words, but for these academic abstracts, 
            # the text is usually uniform. 
            
    # Process tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original_text = p.text
                    if not original_text.strip():
                        continue
                    
                    clean_text = original_text.replace("赛 题", "赛题").replace("场 景", "场景").replace("编排 层", "编排层")
                    polished_text = clean_text
                    for old_str, new_str in REPLACEMENTS.items():
                        clean_old = old_str.replace("赛 题", "赛题").replace("场 景", "场景").replace("编排 层", "编排层")
                        polished_text = polished_text.replace(clean_old, new_str)
                        
                    if polished_text != original_text:
                        p.text = polished_text

    doc.save(output_path)
    logging.info(f"Successfully processed and saved to {output_path}")

if __name__ == "__main__":
    file1 = "E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_论文标准版_修订版.docx"
    file2 = "E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_论文标准版_修订版.docx"
    
    out1 = "E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_润色版.docx"
    out2 = "E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_润色版.docx"
    
    process_docx(file1, out1)
    process_docx(file2, out2)
