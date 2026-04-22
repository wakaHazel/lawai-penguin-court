import os
import docx

def refactor_to_startup_pitch(file_path, out_path):
    print(f"--- 转换为科创展示版 {os.path.basename(file_path)} ---")
    doc = docx.Document(file_path)
    
    replacements = {
        # 核心思想：这是一份向评委展示项目的客观文件。
        # 不能提“为了迎合评委”、“为了准备答辩”、“因为初赛时间紧”等心路历程或主观动机。
        # 朋友直接删词导致语病，我们这里将其替换为流畅、客观的产品或技术术语。
        
        # 修复朋友删词导致的语病：
        "这样的优先级判断，并不是对功能的删减，而是对和项目主轴的尊重。": "这样的优先级判断，并非对功能的简单删减，而是基于产品核心逻辑的聚焦。",
        "对于，首先关心的是：": "在现阶段，本项目的核心目标是验证：",
        "被法学生重复训练、的交互式系统。": "被法学生重复训练、具有实际应用价值的交互式系统。",
        "在价值层面，该系统提供了一种较为完整的智慧法律系统解决方案表达路径，能够更好体现对于“法律需求—技术方案—工作流设计—演示闭环”的要求。": "在技术价值层面，该系统提供了一种较为完整的智慧法律系统解决方案，实现了从“法律需求、技术方案”到“工作流落地”的完整闭环。",
        "可用于。": "具备了向实际应用场景推广的条件。",
        "（）典型使用流程示意": "（四）典型使用流程示意",
        "本节将从工具职责、工作流主链路、节点设计和 Prompt 设计四个层面，说明项目如何使用这些 AI 工具来构建系统。": "本节将从模块职责、工作流主链路、节点设计和 Prompt 设计四个层面，阐述系统底层的AI构建逻辑。",

        # 进一步清理残留的、或需要重写的“心路历程”和比赛痕迹
        "对于区域初赛来说，评委首先关心的是": "在现阶段，本项目的核心目标是验证",
        "为了避免在材料中出现“工具写了很多，但分工不清、工作流不明、Prompt 设计模糊”的问题": "为了清晰呈现技术架构",
        "赛题本身已经明确指出": "行业发展趋势表明",
        "赛题本身已经明确指出：": "行业发展趋势表明：",
        "根据赛题要求": "基于项目规划",
        "赛题工具链约束": "现有技术工具约束",
        "比赛周期": "研发周期",
        "比赛展示场景适配能力": "快速产品演示能力",
        "评委通常会在较短时间内判断": "行业专家或用户需要在较短时间内判断",
        "向评委展示项目闭环与创新点": "清晰呈现项目闭环与技术创新",
        "被评委直观看到成果": "具有实际应用价值",
        "区域初赛阶段的正式申报与后续答辩表达": "具备了向实际应用场景推广的条件",
        "对于本次区域初赛及后续答辩而言": "对于系统的初步验证而言",
        "对于比赛场景来说，这种可讲述、可演示、可证明成果的能力，是作品竞争力的重要组成部分。": "这种从输入到输出的完整闭环演示，是验证产品可行性的关键。",
        
        # 补充材料文档的清理
        "区域初赛Demo与补充材料研究文稿": "项目系统原型与技术补充说明",
        "作为区域初赛的配套说明": "作为本项目的技术补充文档",
        "项目演示方案、补充材料结构与答辩支撑内容": "系统原型演示方案与技术补充材料结构",
        "支撑评委对项目技术逻辑与应用价值的理解": "全面阐释项目的技术逻辑与应用价值",
        "在评审中的定位": "在项目验证中的定位",
        "在初赛评审中": "在系统验证中",
        "项目在评审环节形成了": "项目在此阶段形成了",
        "若比赛阶段采用": "若当前阶段采用",
        "初赛的演示案例": "原型的演示案例",
        "对法律 AI 赛题而言 ，评审通常重点关注": "对于法律 AI 系统的评估，重点通常在于",
        "对法律 AI 赛题而言，评审通常重点关注": "对于法律 AI 系统的评估，重点通常在于",
        "有助于评委确认项目并非停留在概念层": "有助于证明项目并非停留在概念层",
        "对应赛题对工具、工作流和 Prompt 的要求": "对应了系统底层对工具、工作流和 Prompt 的设计要求",
        "该材料使评委在不亲自点击系统的情况下": "该材料能够在非在线演示的情况下",
        "答辩准备": "成果发布准备",
        "等评审关注问题": "等关键技术问题",
        "答辩支撑；提交设计": "系统设计；应用落地"
    }

    # 要整段删除的段落关键字（彻底删除“比赛展示与答辩场景”等不该出现在给评委看的文件里的内容）
    paragraphs_to_remove = [
        "比赛展示与答辩场景",
        "区域初赛材料映射表",
        "表A-1",
        "本附录列示区域初赛",
        "评委通常会在较短时间内判断一个作品",
        "“企鹅法庭”的展示方式非常适合这一场景",
        "这部分要求说明"
    ]

    for p in doc.paragraphs:
        # 1. 检查是否需要整段删除
        should_remove = False
        for kw in paragraphs_to_remove:
            if kw in p.text:
                should_remove = True
                break
        
        if should_remove:
            p.clear()
            continue

        # 2. 文本替换
        original_text = p.text
        new_text = original_text
        for old_str, new_str in replacements.items():
            if old_str in new_text:
                new_text = new_text.replace(old_str, new_str)
        
        if original_text != new_text:
            print(f"修改前: {original_text}")
            print(f"修改后: {new_text}\n")
            p.clear()
            p.add_run(new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original_text = p.text
                    new_text = original_text
                    for old_str, new_str in replacements.items():
                        if old_str in new_text:
                            new_text = new_text.replace(old_str, new_str)
                    if original_text != new_text:
                        p.clear()
                        p.add_run(new_text)

    doc.save(out_path)
    print(f"✅ 保存到: {out_path}\n")

files = [
    r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_流畅严谨版.docx",
    r"E:\lawai\output\doc\区域初赛提交材料_项目Demo与补充材料_企鹅法庭_流畅严谨版.docx"
]

for f in files:
    out_path = f.replace("流畅严谨版", "科创展示版_最终定稿")
    refactor_to_startup_pitch(f, out_path)
