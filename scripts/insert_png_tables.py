import os
import docx
from docx.shared import Inches

def replace_tables_with_png(file_path, out_path, prefix):
    print(f"--- 替换 {os.path.basename(file_path)} 的表格为 PNG ---")
    doc = docx.Document(file_path)
    
    for i, table in enumerate(doc.tables):
        png_filename = f"{prefix}_table_{i+1}.png"
        png_path = os.path.join(r"E:\lawai\output\doc\diagrams\tables", png_filename)
        
        p = table._element.getprevious()
        if p is not None and p.tag.endswith('p'):
            para = docx.text.paragraph.Paragraph(p, doc._body)
            if os.path.exists(png_path):
                # 插入图片，宽度设置为 6 英寸（适应A4纸边距）
                run = para.add_run()
                run.add_picture(png_path, width=Inches(6.0))
            else:
                para.add_run(f"\n[图片未找到: {png_filename}]")
            
        # 物理删除原表格的 XML 节点
        table._element.getparent().remove(table._element)

    doc.save(out_path)
    print(f"✅ 保存到: {out_path}")

files_to_process = {
    r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_科创展示版_最终定稿.docx": "intro",
    r"E:\lawai\output\doc\区域初赛提交材料_项目Demo与补充材料_企鹅法庭_科创展示版_最终定稿.docx": "demo"
}

for f_path, prefix in files_to_process.items():
    out_path = f_path.replace("最终定稿", "图表增强版")
    replace_tables_with_png(f_path, out_path, prefix)
