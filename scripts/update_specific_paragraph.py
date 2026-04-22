import os
from docx import Document

def find_and_replace_text(doc_path, old_text_snippet, new_text):
    if not os.path.exists(doc_path):
        print(f"File not found: {doc_path}")
        return False

    doc = Document(doc_path)
    modified = False
    
    # 因为这段话在之前的润色中可能存在细微的格式变化，我们用关键字匹配来找段落
    keywords = ["优先级判断", "不是对功能的删减", "比赛周期", "项目主轴的尊重", "帮助用户完成一次可感知"]
    
    for p in doc.paragraphs:
        # 如果段落中包含大多数关键字，说明就是这一段
        matches = sum(1 for kw in keywords if kw in p.text)
        if matches >= 3:
            print(f"Found paragraph to modify: \n{p.text}\n")
            p.clear()
            p.add_run(new_text)
            modified = True
            break
            
    if modified:
        doc.save(doc_path)
        print(f"✅ Successfully updated {os.path.basename(doc_path)}")
        return True
    else:
        print(f"⚠️ Could not find the target paragraph in {os.path.basename(doc_path)}")
        return False

def main():
    # 待修改的新文本：去掉了“预判评委”的主观感，改为客观的研发策略和比赛考核点对齐。
    new_text = "这种优先级界定并非对系统功能的简单删减，而是基于项目研发周期与核心目标的理性聚焦。在区域初赛的评审考核维度中，项目的核心价值在于：是否精准定位真实痛点、是否形成完整的业务闭环、是否深度融合AI工具链，以及工程完成度是否达标。基于上述导向，“企鹅法庭”将研发资源高度收敛于“帮助用户完成一次可感知、可推演、可复盘的庭前准备体验”这一主线上，从而确保项目在价值论证与成果交付时更具针对性与说服力。"

    # 这个段落大概率在《项目简介》文档中
    intro_doc = "E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_图文并茂版.docx"
    demo_doc = "E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_图文并茂版.docx"
    
    find_and_replace_text(intro_doc, "", new_text)
    find_and_replace_text(demo_doc, "", new_text)

if __name__ == "__main__":
    main()
