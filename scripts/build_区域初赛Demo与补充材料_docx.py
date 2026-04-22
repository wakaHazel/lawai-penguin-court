from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
DEFAULT_OUTPUT_PATH = OUTPUT_DIR / "区域初赛提交材料_项目Demo与补充材料_企鹅法庭_论文标准版.docx"
OUTPUT_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUTPUT_PATH


def set_run_font(run, east_asia: str, size: int, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)


def clear_paragraph(paragraph) -> None:
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False, size: int = 10, font: str = "宋体") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title_paragraph(doc: Document, text: str, size: int, *, bold: bool = True, spacing_after: int = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    if level == 1:
        paragraph.style = doc.styles["Heading 1"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
    else:
        paragraph.style = doc.styles["Heading 2"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", 16 if level == 1 else 14, bold=True)


def add_body(doc: Document, text: str, *, indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.85)
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12)


def add_list_item(doc: Document, prefix: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.hanging_indent = Cm(0.5)
    paragraph.paragraph_format.space_after = Pt(4)
    run_prefix = paragraph.add_run(prefix)
    set_run_font(run_prefix, "宋体", 12, bold=True)
    run_text = paragraph.add_run(text)
    set_run_font(run_text, "宋体", 12)


def add_catalog_entry(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    set_run_font(run, "宋体", 12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True, center=True, font="黑体")
        clear_cell_shading(header_cells[idx])
        if col_widths_cm:
            header_cells[idx].width = Cm(col_widths_cm[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=10)
            if col_widths_cm:
                cells[idx].width = Cm(col_widths_cm[idx])
    doc.add_paragraph()


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def configure_section_layout(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.6)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)
    section.different_first_page_header_footer = False


def _append_field_run(paragraph, instr: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run_begin = paragraph.add_run()
    run_begin._r.append(begin)
    run_instr = paragraph.add_run()
    run_instr._r.append(instr_text)
    run_sep = paragraph.add_run()
    run_sep._r.append(separate)
    run_text = paragraph.add_run("1")
    run_text._r.append(end)
    set_run_font(run_text, "宋体", 9, color=RGBColor(0x66, 0x66, 0x66))


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_section_header(section, text: str | None = None) -> None:
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, "宋体", 10, color=RGBColor(0x44, 0x44, 0x44))


def set_section_footer(section, *, roman: bool = False, visible: bool = True) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not visible:
        return
    prefix = paragraph.add_run("— ")
    set_run_font(prefix, "宋体", 9, color=RGBColor(0x66, 0x66, 0x66))
    _append_field_run(paragraph, "PAGE \\* ROMAN" if roman else "PAGE")
    suffix = paragraph.add_run(" —")
    set_run_font(suffix, "宋体", 9, color=RGBColor(0x66, 0x66, 0x66))


def add_cover_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, center=True, size=11, font="黑体")
        set_cell_text(cells[1], value, size=11)
        clear_cell_shading(cells[0])
        clear_cell_shading(cells[1])
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(10.4)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    configure_section_layout(section)

    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(12)
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    set_section_header(section, None)
    set_section_footer(section, visible=False)


def build_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    add_title_paragraph(doc, "【D06】法律AI应用创新与实践", 18, spacing_after=18)
    add_title_paragraph(doc, "企鹅法庭·沉浸式庭审模拟与法律辅助分析系统", 24, spacing_after=12)
    add_title_paragraph(doc, "区域初赛Demo与补充材料研究文稿", 16, spacing_after=16)
    add_title_paragraph(doc, "项目演示方案、补充材料结构与答辩支撑内容", 13, bold=False, spacing_after=28)
    for _ in range(3):
        doc.add_paragraph()
    add_cover_meta_table(
        doc,
        [
            ("赛题名称", "【D06】法律AI应用创新与实践"),
            ("文稿名称", "区域初赛Demo与补充材料研究文稿"),
            ("项目名称", "企鹅法庭·沉浸式庭审模拟与法律辅助分析系统"),
            ("适用阶段", "区域初赛"),
            ("项目团队", "企鹅法庭项目组"),
            ("完成时间", "2026年4月"),
        ],
    )
    add_title_paragraph(doc, "文稿属性：正式配套研究文稿", 12, bold=False, spacing_after=8)


def build_catalog(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "目  录", level=1)
    items = [
        "摘  要",
        "第一章  项目Demo研究说明",
        "第二章  补充材料组织与提交设计",
    ]
    for item in items:
        add_catalog_entry(doc, item)


def build_abstract(doc: Document) -> None:
    add_heading(doc, "摘  要", level=1)
    add_body(
        doc,
        "作为区域初赛配套文稿，本研究文稿围绕“企鹅法庭·沉浸式庭审模拟与法律辅助分析系统”的项目演示方案与补充材料组织方式展开，重点说明作品 Demo 的展示主线、视频脚本、案例选取原则、提交项构成，以及补充材料如何服务于项目完成度、专业性和答辩说服力的表达。文稿遵循“展示主轴清晰、证据材料可核、系统能力可映射、评审问题可直接响应”的组织原则，将项目 Demo 与补充材料从单纯的提交附件提升为支撑作品理解、作品验证与现场答辩的正式配套研究材料。",
        indent=False,
    )
    add_body(
        doc,
        "文稿首先分析作品 Demo 在区域初赛中的功能定位，明确其不仅承担展示系统界面的作用，更承担验证系统完成度、解释主工作流、证明工具链落地与引导演示节奏的任务；继而从项目技术方案、AI 工具与工作流说明、输入输出示例包、团队分工与专业验证说明、部署与运行说明等方面，构建补充材料的总体结构。通过上述设计，项目可以在正式提交与答辩场景中形成“主报告负责问题与方案、Demo负责系统展示、补充材料负责证据支撑”的三层表达体系，从而提高区域初赛材料的整体完整性与学术规范性。",
        indent=False,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run_label = paragraph.add_run("关键词：")
    set_run_font(run_label, "黑体", 12, bold=True)
    run_text = paragraph.add_run("项目Demo；补充材料；智慧法律；庭审模拟；答辩支撑；提交设计")
    set_run_font(run_text, "宋体", 12)


def build_demo_section(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "第一章  项目Demo研究说明", level=1)
    add_body(
        doc,
        "项目 Demo 需要说明作品当前的呈现形态、核心能力、演示过程中展示的内容，以及与“项目简介”中功能说明的对应关系。对区域初赛而言，作品 Demo 是评委判断项目完成度和可运行性的直接依据，因此材料同时覆盖“可访问、可理解、可复现、可答辩”四个维度。",
    )

    add_heading(doc, "（一）作品Demo形态说明", level=2)
    add_body(
        doc,
        "“企鹅法庭”当前作品形态为 Web 端应用系统，整体围绕“新建案件—进入模拟—推进庭审—查看分析—输出复盘”这一主流程展开。作品前端提供案件录入、历史案件查看、庭审模拟互动、对方行为推进、胜诉率分析和复盘报告展示等界面；后端提供案件会话管理、历史轨迹记录、检查点恢复、结果分析与报告生成能力；同时结合腾讯元器工作流与法律检索能力，形成一套可演示、可扩展的法律 AI 系统原型。",
    )
    add_body(
        doc,
        "项目 Demo 采用“预置 Demo 案件 + 实时互动 + 结果落地”的演示方式。演示开场直接选取系统预设的典型案件模板，快速进入案件录入后的主工作台，再通过一轮到两轮关键节点互动，展示庭审文本如何推进、对方行为如何生成、用户策略如何影响后续变化，最后进入结果分析与复盘报告页面，完整呈现项目闭环。",
    )

    add_heading(doc, "（二）作品Demo提交信息", level=2)
    add_table(
        doc,
        ["Demo字段", "提交内容", "备注"],
        [
            ["作品名称", "企鹅法庭·沉浸式庭审模拟与法律辅助分析系统", "与项目简介、视频封面和答辩 PPT 保持同名"],
            ["作品形态", "Web 应用 / 智能体系统解决方案", "以网页交互界面为主体，结合工作流与后端控制层实现"],
            ["演示入口", "正式公网地址（提交前填入）", "如完成正式部署，填入 Render 或正式演示域名"],
            ["代码仓库", "https://github.com/wakaHazel/lawai-penguin-court", "作为补充材料，不作为正式作品入口"],
            ["本地演示启动方式", "运行“一键启动完整服务.bat”", "适用于现场本地稳定演示与录制视频"],
            ["主要演示链路", "案件录入 -> 庭审模拟 -> 胜诉率分析 -> 复盘报告", "作为视频和现场答辩固定主线"],
        ],
        col_widths_cm=[3.0, 6.2, 5.0],
    )
    add_body(
        doc,
        "作品正式公网链接在提交前统一替换为最终可访问的演示地址。若比赛阶段采用本地演示而不提供外网完整服务，可在材料中写明“系统提供本地稳定演示版本，并附视频演示与界面展示”，同时提交源码仓库、系统截图和视频文件作为辅助证明材料。",
    )

    add_heading(doc, "（三）Demo核心演示内容", level=2)
    add_body(
        doc,
        "项目 Demo 围绕最有代表性的主链路展开，依次展示：首页的“新建案件/历史案件”入口、典型案件模板选择、案件录入与推演工作台、庭审互动关键节点，以及胜诉率分析和复盘报告页面。整条链路对应案件录入、阶段推进、对方模拟、结果分析和成果沉淀五个核心环节。",
    )
    add_body(
        doc,
        "该演示链路能够在有限时间内集中呈现案件录入、流程推进、结果分析与报告沉淀四类核心能力，从而较清晰地区分本项目与单轮法律问答型产品在产品形态和使用方式上的差异。若时间允许，可在主线演示之外补充系统还支持历史案件续接、回放检查点、复盘报告查看和多类预置案件模板等能力。",
    )

    add_heading(doc, "（四）Demo视频脚本", level=2)
    add_table(
        doc,
        ["镜头顺序", "画面内容", "讲解重点", "参考时长"],
        [
            ["镜头1", "项目标题页与系统首页", "说明项目定位：不是泛法务问答，而是诉前准备与庭审模拟系统", "20-30秒"],
            ["镜头2", "选择典型 Demo 案件模板", "说明案件类型、适用人群和本次演示目标", "20秒"],
            ["镜头3", "进入案件录入与主工作台", "展示结构化案情输入、争点识别与阶段主轴", "30-40秒"],
            ["镜头4", "推进庭审模拟关键节点", "展示法官提示、对方行为、用户分支动作和剧情继续展开", "60-90秒"],
            ["镜头5", "查看胜诉率分析与复盘报告", "展示风险点、证据缺口、准备建议和报告价值", "40-60秒"],
            ["镜头6", "收尾总结页", "总结项目价值、AI工具链使用和应用前景", "20-30秒"],
        ],
        col_widths_cm=[1.8, 4.0, 6.0, 3.2],
    )
    add_body(
        doc,
        "视频总时长控制在 3 至 5 分钟。旁白内容围绕三个问题展开：系统解决什么问题、系统与普通法律 AI 的区别、系统如何依托腾讯系 AI 工具链实现。若条件允许，可穿插一到两张系统原型界面或关键结果页截图，用于强化界面完整度和视觉表达效果。",
    )

    add_heading(doc, "（五）Demo演示案例", level=2)
    add_body(
        doc,
        "正式演示样本从系统现有的预置案例中选取。当前项目已围绕劳动争议、民间借贷、离婚纠纷和侵权责任等方向整理多组 Demo 案件模板，其中劳动争议和民间借贷两类案件更适合区域初赛演示。前者能够体现劳动关系认定、考勤与工资证据、双倍工资请求等问题，后者则能更直观地展示借贷事实、转账凭证、聊天记录和抗辩路径。",
    )
    add_body(
        doc,
        "固定演示 1 到 2 个高质量案例，更能体现项目是否真正把一个案例走完整，而不是只展示多个案由选项。其他案例模板可在补充材料中以“系统已支持或已预置案例库”形式展示，用于体现项目扩展基础。",
    )

    add_heading(doc, "（六）Demo提交材料清单", level=2)
    add_table(
        doc,
        ["提交项", "提交内容", "说明"],
        [
            ["作品Demo链接", "正式公网演示地址（待填）", "若无稳定公网地址，可在说明中补充本地演示版本与视频"],
            ["作品界面截图", "首页、案件录入页、庭审模拟页、分析报告页", "用于增强材料的可视化说服力"],
            ["演示视频", "MP4 文件，3-5 分钟", "包含操作过程、界面展示与旁白说明"],
            ["视频脚本", "演示视频文稿或口播脚本", "便于后续彩排与答辩统一口径"],
            ["源码仓库链接", "GitHub 仓库链接", "作为补充佐证材料"],
        ],
        col_widths_cm=[3.0, 5.0, 6.2],
    )


def build_supplementary_section(doc: Document) -> None:
    add_page_break(doc)
    add_heading(doc, "第二章  补充材料组织与提交设计", level=1)
    add_body(
        doc,
        "补充材料直接服务于项目完成度、专业性和可信度的展示。对法律 AI 赛题而言，评审通常重点关注法律内容支撑、工作流设计、人工审校与风险控制、输入输出样例完整度。补充材料围绕这些重点展开。",
    )

    add_heading(doc, "（一）补充材料类型", level=2)
    add_table(
        doc,
        ["材料类别", "材料内容", "对应价值"],
        [
            ["项目技术方案", "系统架构、前后端分工、元器工作流与后端协同说明", "证明项目不是停留在概念层，而是有清晰实现路径"],
            ["AI工具使用说明", "工具清单、工作流链路、Prompt 设计、风险控制", "直接回应赛题对工作流与 Prompt 的要求"],
            ["输入输出示例包", "典型案件输入、模拟结果、复盘报告输出", "证明系统具备真实可读结果"],
            ["系统截图或原型图", "首页、案件录入、庭审模拟、报告页", "增强视觉完成度和产品感"],
            ["团队分工与验证说明", "五人分工、法学审校机制、专业验证路径", "增强专业可信度"],
            ["GitHub仓库或部署说明", "代码仓库地址、启动方式、部署方案", "体现工程实现能力与可复现性"],
        ],
        col_widths_cm=[3.2, 6.0, 5.0],
    )

    add_heading(doc, "（二）补充材料一：系统技术方案", level=2)
    add_body(
        doc,
        "系统技术方案用于说明系统架构、模块边界、数据流和主工作流设计，直接回答系统如何组织、官方工具链承担何种角色、元器与后端如何协同、前端体验是否与工作流联动等问题。",
    )
    add_body(
        doc,
        "当前仓库中已有较成熟的技术方案和相关设计文稿，可直接整理后作为补充材料底稿使用。该材料有助于评委确认项目并非停留在概念层，而是具备结构明确、实现边界清晰、阶段目标明确的工程基础。",
    )

    add_heading(doc, "（三）补充材料二：AI工具与工作流说明", level=2)
    add_body(
        doc,
        "AI 工具与工作流说明对应赛题对工具、工作流和 Prompt 的要求。材料内容包括腾讯开悟比赛平台、CodeBuddy、腾讯元器、腾讯元宝、小理AI与得理开放平台API等工具在项目中的职责边界与协同关系。",
    )
    add_body(
        doc,
        "材料中还可简要展示元器工作流主链路图、关键节点说明表和三类核心 Prompt 的设计逻辑。该材料能够直接回应工作流如何落地、官方工具如何实际使用等评审关注问题，并增强项目在技术实现层面的完整度。",
    )

    add_heading(doc, "（四）补充材料三：输入输出示例包", level=2)
    add_body(
        doc,
        "输入输出示例包用于展示系统结果质量。可从系统已经预置的 Demo 案件中选取 1 至 3 组代表性案例，整理为“输入内容 + 系统输出结果”的标准化样例。输入部分包括案件类型、案件摘要、主要诉讼请求、核心证据和对方信息；输出部分包括当前争点、模拟节点结果、风险提示、胜诉率分析摘要和复盘报告节选。",
    )
    add_body(
        doc,
        "该材料使评委在不亲自点击系统的情况下，也能直接判断项目输出是否具备实质内容。样例中若能清楚体现争点分析、证据缺口、对方抗辩路径和后续准备动作，将明显增强作品可信度。",
    )

    add_heading(doc, "（五）补充材料四：团队分工与专业验证说明", level=2)
    add_body(
        doc,
        "团队分工与专业验证说明用于回应法律内容是否经过专业审校、团队是否具备法律与技术协同能力等关键问题。材料中应明确五人团队的职责分工，尤其突出法学同学在争点整理、法条校验、案例筛选、输出审校和风险控制中的作用。",
    )
    add_body(
        doc,
        "材料内容保持简洁清晰即可。例如：技术同学负责前后端实现、工作流联调与部署；法学同学负责法律素材库整理、规则表建立和 AI 输出人工校验；项目负责人负责材料统筹、演示视频和答辩准备。这样的分工说明能够体现项目具备法律内容的专业把关机制。",
    )

    add_heading(doc, "（六）补充材料五：部署与运行说明", level=2)
    add_body(
        doc,
        "部署与运行说明列明系统通过何种方式启动与演示，例如本地完整服务启动脚本、前后端独立启动方式、可选公网部署路径等。该材料能够说明作品并非一次性演示稿，而是具备重复运行和继续完善基础的系统原型。",
    )
    add_body(
        doc,
        "目前项目已具备本地完整服务运行能力，并已有公开代码仓库可供补充说明。在正式提交中，可以把“本地演示版本稳定可运行”和“提供代码仓库与部署说明”作为工程完成度的佐证。即使公网服务尚未最终稳定，也可以通过这份材料表明：系统具备可复现的交付基础，而不是只存在于 PPT 或视频中。",
    )

    add_heading(doc, "（七）补充材料清单总表", level=2)
    add_table(
        doc,
        ["序号", "补充材料名称", "提交优先级", "备注"],
        [
            ["1", "项目技术方案", "高", "用于说明系统架构、主流程与工程实现路径"],
            ["2", "AI工具使用说明（工作流与Prompt）", "高", "直接响应赛题要求，强化官方工具链使用证明"],
            ["3", "输入输出示例包", "高", "展示系统输出质量与法律场景落地能力"],
            ["4", "系统界面截图/原型图", "高", "增强视觉说服力与完成度"],
            ["5", "团队分工与专业验证说明", "高", "说明法律审校机制与团队协作结构"],
            ["6", "GitHub仓库与部署说明", "中", "作为工程实现与可复现性的补充佐证"],
            ["7", "答辩问答库或演示脚本", "可选", "适合作为内部准备材料，也可酌情作为补充附件"],
        ],
        col_widths_cm=[1.2, 5.2, 2.4, 7.2],
    )
    add_body(
        doc,
        "补充材料选取原则为“少而精、强关联、能证明”。提交内容围绕系统是否完成、官方工具是否实际使用、输出是否具有实质内容、法律内容是否经过专业校验四个方面组织。",
    )


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    build_cover(doc)

    front_matter = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section_layout(front_matter)
    set_page_number_start(front_matter, 1)
    set_section_header(front_matter, None)
    set_section_footer(front_matter, roman=True, visible=True)
    build_abstract(doc)
    build_catalog(doc)

    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section_layout(main_section)
    set_page_number_start(main_section, 1)
    set_section_header(main_section, "企鹅法庭·区域初赛Demo与补充材料研究文稿")
    set_section_footer(main_section, roman=False, visible=True)
    build_demo_section(doc)
    build_supplementary_section(doc)
    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
