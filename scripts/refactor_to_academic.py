import os
import docx

def refactor_to_academic(file_path, out_path):
    print(f"--- 转换为学术风格 {os.path.basename(file_path)} ---")
    doc = docx.Document(file_path)
    
    replacements = {
        # 修正商业化词汇
        "商业验证": "实践检验",
        "商业落地": "实际应用",
        "商业白皮书": "项目研究报告",
        "商业化落地": "应用落地",
        "商业潜创团队": "科研团队",
        "市场和资方": "学术界与实务界",
        "产品 MVP (最小可行性产品)": "系统原型 (Prototype)",
        "MVP (最小可行性产品)": "系统原型",
        "MVP版本": "原型系统",
        "MVP": "系统原型",
        "初版MVP": "第一版系统原型",
        "产品研发周期": "课题研究周期",
        "产品发布": "成果发布",
        "产品验收": "项目结题与验收",
        "成熟的法律科技产品": "成熟的法律AI系统",
        "一款成熟的法律科技产品": "一个完善的法律AI科研项目",
        "产品展示": "成果展示",
        "产品操作视频": "系统运行视频",
        "业务逻辑": "技术逻辑",
        "业务链路": "功能链路",
        "核心业务工作流图解": "核心功能工作流图解",
        "核心业务链路": "核心技术链路",
        "主业务流": "主控工作流",
        "业务延展性": "系统扩展性",
        "行业专家": "领域专家",
        "行业对": "实务界对",
        "技术白皮书": "技术研究报告",
        "客户": "用户",
        "投资人": "领域专家",
        
        # 将部分过于直白的验证词汇改回更学术严谨的表述
        "功能验证方案": "系统测试与验证方案",
        "验证样本": "测试样本",
        "验证案例": "测试案例",
        "标准验证案例": "标准测试案例",
        "验证过程": "测试过程",
        "重点验证": "重点测试"
    }

    for p in doc.paragraphs:
        original_text = p.text
        new_text = original_text
        for old_str, new_str in replacements.items():
            if old_str in new_text:
                new_text = new_text.replace(old_str, new_str)
        
        if original_text != new_text:
            print(f"修改前: {original_text}")
            print(f"修改后: {new_text}\n")
            p.clear()
            p.add_run(new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original_text = p.text
                    new_text = original_text
                    for old_str, new_str in replacements.items():
                        if old_str in new_text:
                            new_text = new_text.replace(old_str, new_str)
                    if original_text != new_text:
                        p.clear()
                        p.add_run(new_text)

    doc.save(out_path)
    print(f"✅ 保存到: {out_path}\n")

files = [
    r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_商业白皮书_最终版.docx",
    r"E:\lawai\output\doc\区域初赛提交材料_项目Demo与补充材料_企鹅法庭_商业白皮书_最终版.docx"
]

for f in files:
    out_path = f.replace("商业白皮书_最终版", "学术科创版_最终版")
    refactor_to_academic(f, out_path)
