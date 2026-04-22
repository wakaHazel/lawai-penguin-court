import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

def insert_image_after_text(doc, search_text, image_path, caption_text):
    for idx, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            if idx + 2 < len(doc.paragraphs) and caption_text in doc.paragraphs[idx+2].text:
                img_p = doc.paragraphs[idx + 1]
                img_p.clear()
                run = img_p.add_run()
                try:
                    run.add_picture(image_path, width=Inches(6.0)) # 稍微大一点，架构图需要清晰
                    print(f"✅ Updated existing logic diagram {image_path} for caption '{caption_text}'")
                except Exception as e:
                    print(f"Failed to update logic diagram {image_path}: {e}")
                return
            
            new_p = OxmlElement('w:p')
            p._p.addnext(new_p)
            img_p = doc.paragraphs[idx + 1]
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = img_p.add_run()
            try:
                run.add_picture(image_path, width=Inches(6.0))
            except Exception as e:
                print(f"Failed to insert logic diagram {image_path}: {e}")
                return
                
            new_caption_p = OxmlElement('w:p')
            img_p._p.addnext(new_caption_p)
            caption_p = doc.paragraphs[idx + 2]
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            caption_run = caption_p.add_run(caption_text)
            caption_run.font.size = Pt(10)
            caption_run.font.bold = True
            caption_p.paragraph_format.space_after = Pt(12)
            
            print(f"✅ Inserted new logic diagram {image_path} with caption '{caption_text}'")
            return
            
    print(f"⚠️ Could not find text '{search_text}' in document.")

def main():
    base_diagram_dir = "E:/lawai/output/doc/diagrams"
    img_arch = os.path.join(base_diagram_dir, "system_architecture.png")
    img_workflow = os.path.join(base_diagram_dir, "core_workflow.png")
    
    # 我们将把架构图和流程图插入到《项目简介》文档的核心技术方案章节
    intro_doc_path = "E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_最终提交版.docx"
    intro_out_path = "E:/lawai/output/doc/区域初赛提交材料_项目简介_企鹅法庭_架构图文版.docx"
    
    if os.path.exists(intro_doc_path):
        doc = Document(intro_doc_path)
        
        # 1. 插入四层架构图
        insert_image_after_text(doc, "在架构设计上，系统采用了“前端交互、后端控制、腾讯元器编排、法律能力支撑”的四层模块化结构：", img_arch, "图 1：企鹅法庭·系统总体技术架构图")
        
        # 2. 插入核心工作流图
        insert_image_after_text(doc, "系统以民事诉讼庭前准备阶段为核心，建立了“案件录入—庭审推演—对手模拟—胜率评估—复盘报告”的完整工作流", img_workflow, "图 2：企鹅法庭·核心业务流转与多阶段状态机设计")
        
        doc.save(intro_out_path)
        print(f"Saved: {intro_out_path}")
        
    # Demo 文档也插一张主干流程图，用于解释操作
    demo_doc_path = "E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_最终提交版.docx"
    demo_out_path = "E:/lawai/output/doc/区域初赛提交材料_项目Demo与补充材料_企鹅法庭_架构图文版.docx"
    
    if os.path.exists(demo_doc_path):
        doc = Document(demo_doc_path)
        insert_image_after_text(doc, "目前，“企鹅法庭”已开发为完整的 Web 端应用系统。其业务主线遵循“案件录入—庭审推演—互动质证—胜率分析—复盘输出”的逻辑顺序", img_workflow, "图 1：Demo 演示主线——核心业务工作流图解")
        doc.save(demo_out_path)
        print(f"Saved: {demo_out_path}")

if __name__ == "__main__":
    main()
