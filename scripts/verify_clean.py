import os
import docx

def verify_clean(file_path):
    print(f"--- 验证 {os.path.basename(file_path)} ---")
    doc = docx.Document(file_path)
    keywords = ["答辩", "评审", "评委", "比赛", "初赛", "参赛", "区域赛"]
    found = False
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        for kw in keywords:
            if kw in text:
                print(f"⚠️ 警告！发现残留关键字 '{kw}' 在段落 {i}: {text}")
                found = True
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for kw in keywords:
                        if kw in text:
                            print(f"⚠️ 警告！发现残留关键字 '{kw}' 在表格中: {text}")
                            found = True

    if not found:
        print("✅ 验证通过：文档非常纯净，没有比赛相关字眼！")

files = [
    r"E:\lawai\output\doc\区域初赛提交材料_项目简介_企鹅法庭_去比赛痕迹_纯净版.docx",
    r"E:\lawai\output\doc\区域初赛提交材料_项目Demo与补充材料_企鹅法庭_去比赛痕迹_纯净版.docx"
]

for f in files:
    verify_clean(f)
